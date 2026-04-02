"""
export/docx_template_export.py
================================
Renders the Jinja2 DOCX template (`docs_template/TP_Local_File_TEMPLATE.docx`)
using **docxtpl** and the current Streamlit session state.

Entry point:
    render_tp_document(state: dict, output_path: str) -> None

The `state` dict is built from `st.session_state` in `app.py`.
"""

from __future__ import annotations

import os
import io
import tempfile
from pathlib import Path
from typing import Any

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches

# ─────────────────────────────────────────────────────────────────────────────
# Template path (relative to repo root / tp_app/)
# ─────────────────────────────────────────────────────────────────────────────
_HERE = Path(__file__).parent.parent          # tp_app/
TEMPLATE_PATH = _HERE / "TP_Local_File_TEMPLATE.docx"


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _fmt_idr(value: Any) -> str:
    """Format a number as IDR string, e.g. 1234567 -> 'IDR 1,234,567'."""
    try:
        return f"IDR {int(float(str(value).replace(',', '').replace('IDR', '').strip())):,}"
    except Exception:
        return str(value)


def _pct(value: Any) -> str:
    """Return value as a percentage string, e.g. 0.113 -> '11.3' or '11.30%' -> '11.30'."""
    try:
        s = str(value).replace('%', '').strip()
        return f"{float(s):.2f}"
    except Exception:
        return str(value)


def _safe(value: Any, default: str = "") -> str:
    """Return string value, or default if empty/None."""
    if value is None:
        return default
    s = str(value).strip()
    return s if s else default


def _strip_md(text: str) -> str:
    """
    Strip markdown and fix paragraph rendering for docxtpl.
    - Removes ## heading markers, **bold**, *italic*, hr lines
    - Removes numbered ALL-CAPS heading lines (e.g. '1. PURPOSE OF REPORT')
      which are redundant with Word headings and stretch badly when justified
    - Normalizes bullet markers (* / •) to '- ' and keeps them so that
      _split_content() can apply 'List Paragraph' Word style per bullet
    - Replaces double newlines with \\a (docxtpl paragraph separator)
      so each block is a real Word paragraph, not a line-break in one paragraph
    """
    import re
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    # Remove numbered ALL-CAPS headings like "1. PURPOSE OF REPORT"
    text = re.sub(r'^\d+\.\s+[A-Z][A-Z\s,&()\-]+$', '', text, flags=re.MULTILINE)
    # Remove bare ALL-CAPS heading lines (e.g. "SUPPLY CHAIN DESCRIPTION",
    # "FUNCTIONAL ANALYSIS TABLE NARRATIVE") — template already has these headings.
    # Match lines that are 2+ ALL-CAPS words (letters, spaces, allowed punctuation only).
    text = re.sub(r'^[A-Z][A-Z\s,&()\-:]{4,}$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,2}(.*?)_{1,2}', r'\1', text)
    text = re.sub(r'^[-_*]{3,}\s*$', '', text, flags=re.MULTILINE)
    # Normalize bullet markers (* / •) to '- ' — keep prefix so _split_content()
    # can detect bullet lines and apply List Paragraph style in Word.
    text = re.sub(r'^\*\s+', '- ', text, flags=re.MULTILINE)
    text = re.sub(r'^•\s+', '- ', text, flags=re.MULTILINE)
    # Ensure each bullet line starts a new paragraph (insert blank line before it)
    # so it gets converted to \a and becomes a separate Word paragraph.
    text = re.sub(r'([^\n])\n(- )', r'\1\n\n\2', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()
    # Use \a as docxtpl paragraph separator (creates real Word paragraphs)
    text = text.replace('\n\n', '\a')
    # Single newline -> space (inline within same paragraph)
    text = text.replace('\n', ' ')
    return text


# ─────────────────────────────────────────────────────────────────────────────
# Context builder
# ─────────────────────────────────────────────────────────────────────────────

def build_context(state: dict, tpl: DocxTemplate) -> dict:
    """
    Build the Jinja2 context dict from `state` (session state snapshot).
    All keys must match the {{ variable }} names in TP_Local_File_TEMPLATE.docx.
    """
    qr = state.get("quartile_range", {})
    fd = state.get("financial_data", {})
    fd_prior = state.get("financial_data_prior", {})

    # ── company identity ─────────────────────────────────────────────────────
    company_name  = _safe(state.get("company_name"), "PT [Company Name]")
    company_short = _safe(state.get("company_short_name"), "[Short Name]")
    fiscal_year   = _safe(state.get("fiscal_year"), "2024")
    # brand_name: the product/trade brand (e.g. "SANY" for a SANY distributor).
    # Falls back to company_short if not explicitly set.
    brand_name    = _safe(state.get("brand_name"), company_short)

    # company_name_no_pt: sections 6-16 headers have "PT " as a hardcoded
    # separate run, then {{ company_name_no_pt }} for the rest.
    # Use the full name minus "PT " so the header reads "PT [full name]"
    # and right-aligns (shifts left) if the name is long.
    company_name_no_pt = company_name.removeprefix("PT ").strip()

    # ── shareholders ─────────────────────────────────────────────────────────
    # Auto-calculate ownership percentage from shares if not manually provided
    _raw_shareholders = [s for s in state.get("shareholders", []) if s.get("name")]
    _total_shares = 0
    for _sh in _raw_shareholders:
        try:
            _total_shares += float(str(_sh.get("shares", "0")).replace(",", "").strip() or 0)
        except (ValueError, TypeError):
            pass

    def _calc_pct(sh) -> str:
        """Return ownership % from shares/total, fallback to manual entry."""
        manual = _safe(sh.get("percentage"))
        if _total_shares > 0:
            try:
                shares = float(str(sh.get("shares", "0")).replace(",", "").strip() or 0)
                return f"{shares / _total_shares * 100:.2f}%"
            except (ValueError, TypeError):
                pass
        return manual

    shareholders = [
        {
            "no":         str(i + 1),
            "name":       _safe(sh.get("name")),
            "shares":     _safe(sh.get("shares")),
            "capital":    _safe(sh.get("capital")),
            "percentage": _calc_pct(sh),
        }
        for i, sh in enumerate(_raw_shareholders)
    ]

    # ── management ───────────────────────────────────────────────────────────
    management = [
        {
            "position": _safe(m.get("position")),
            "name":     _safe(m.get("name")),
        }
        for m in state.get("management", [])
        if m.get("name")
    ]

    # ── affiliated parties ────────────────────────────────────────────────────
    affiliated_parties = [
        {
            "name":             _safe(a.get("name")),
            "country":          _safe(a.get("country")),
            "affiliation_type": _safe(a.get("relationship")),
            "transaction_type": _safe(a.get("transaction_type")),
        }
        for a in state.get("affiliated_parties", [])
        if a.get("name")
    ]

    # ── products ──────────────────────────────────────────────────────────────
    products = [
        {
            "name":        _safe(p.get("name")),
            "description": _safe(p.get("description")),
        }
        for p in state.get("products", [])
        if p.get("name")
    ]

    # ── competitors ───────────────────────────────────────────────────────────
    competitors = [
        {
            "no":          str(i + 1),
            "name":        _safe(c.get("name")),
            "description": _safe(c.get("description", c.get("ros_data", ""))),
        }
        for i, c in enumerate(c for c in state.get("comparable_companies", []) if c.get("name"))
    ]

    # ── affiliated / independent transactions ─────────────────────────────────
    # The session state stores these in affiliated_parties with transaction_type.
    # We also allow a dedicated list if present.
    affiliated_transactions = state.get("affiliated_transactions", [])
    if not affiliated_transactions:
        # Build from affiliated_parties (fallback)
        affiliated_transactions = [
            {
                "no":               str(i + 1),
                "name":             _safe(a.get("name")),
                "country":          _safe(a.get("country")),
                "affiliation_type": _safe(a.get("relationship")),
                "transaction_type": _safe(a.get("transaction_type")),
                "type_of_product":  "",
                "amount_idr":       "",
                "quantity":         "",
                "price_per_unit":   "",
            }
            for i, a in enumerate(a for a in state.get("affiliated_parties", []) if a.get("name"))
        ]
    else:
        for i, a in enumerate(affiliated_transactions):
            a.setdefault("no", str(i + 1))
            a.setdefault("type_of_product", a.pop("value", ""))
            a.setdefault("amount_idr",      a.pop("note", ""))
            a.setdefault("quantity",        "")
            a.setdefault("price_per_unit",  "")

    independent_transactions = state.get("independent_transactions", [])
    for i, t in enumerate(independent_transactions):
        t.setdefault("no", str(i + 1))

    # ── search criteria ───────────────────────────────────────────────────────
    search_criteria_results = [
        {
            "step":         _safe(r.get("step", r.get("criteria", ""))),
            "criteria":     _safe(r.get("criteria", r.get("description", ""))),
            "result_count": _safe(r.get("result_count", r.get("result", ""))),
        }
        for r in state.get("search_criteria_results", [])
    ]

    # ── rejection matrix ──────────────────────────────────────────────────────
    CHECKMARK = "√"
    rejection_matrix = [
        {
            "name":                            _safe(r.get("name", r.get("company", ""))),
            "limited_financial_statement":     CHECKMARK if r.get("limited_financial_statement") else "",
            "negative_margin":                 CHECKMARK if r.get("negative_margin") else "",
            "consolidated_financial_statement":CHECKMARK if r.get("consolidated_financial_statement") else "",
            "different_main_activity":         CHECKMARK if r.get("different_main_activity") else "",
            "non_comparable_line_of_business": CHECKMARK if r.get("non_comparable_line_of_business") else "",
            "limited_information_website":     CHECKMARK if r.get("limited_information_website") else "",
            "accepted":                        CHECKMARK if r.get("accepted") else "",
        }
        for r in state.get("rejection_matrix", [])
        if r.get("name")
    ]

    # ── comparable companies ──────────────────────────────────────────────────
    comparable_descriptions = state.get("comparable_descriptions", {}) or {}
    comparable_companies = [
        {
            "name":              _safe(c.get("name")),
            "country":           _safe(c.get("country")),
            "business_activity": _safe(c.get("description")),
            "ros_data":          _safe(c.get("ros_data")),
            "ai_description":    _safe(comparable_descriptions.get(c.get("name", ""), c.get("description", ""))),
        }
        for c in state.get("comparable_companies", [])
        if c.get("name")
    ]

    # ── IQR / ratios ──────────────────────────────────────────────────────────
    q1     = _pct(qr.get("q1",     state.get("q1",     0.0)))
    median = _pct(qr.get("median", state.get("median", 0.0)))
    q3     = _pct(qr.get("q3",     state.get("q3",     0.0)))
    tested_party_ratio = _pct(state.get("tested_party_ratio", 0.0))

    # ── P/L table ─────────────────────────────────────────────────────────────
    def fd_row(label: str, key: str, is_pct: bool = False) -> dict:
        cur  = _safe(fd.get(key))
        prior = _safe(fd_prior.get(key))
        # Calculate % change for performance column
        try:
            c = float(str(cur).replace(',', '').replace('IDR', '').strip())
            p = float(str(prior).replace(',', '').replace('IDR', '').strip())
            perf = f"{((c - p) / abs(p) * 100):.2f}%" if p != 0 else "N/A"
        except Exception:
            perf = ""
        return {
            "account":    label,
            "fy_current": cur,
            "fy_prior":   prior,
            "performance": perf,
        }

    pl_table = [
        fd_row("Sales",                 "sales"),
        fd_row("Cost of Goods Sold",    "cogs"),
        fd_row("Gross Profit",          "gross_profit"),
        fd_row("Operating Expenses",    "operating_expenses"),
        fd_row("Operating Income",      "operating_profit"),
        fd_row("Financial Income",      "financial_income"),
        fd_row("Other Expense",         "other_expense"),
        fd_row("Income Before Tax",     "income_before_tax"),
        fd_row("Income Tax",            "income_tax"),
        fd_row("Net Income",            "net_income"),
    ]
    # Remove rows where both current and prior are empty
    pl_table = [r for r in pl_table if r["fy_current"] or r["fy_prior"]]

    # ── financial helpers for ROS table ───────────────────────────────────────
    fd_prior2 = state.get("financial_data_prior2", {})

    def _ros(fd_dict: dict) -> str:
        try:
            op = float(str(fd_dict.get("operating_profit", 0)).replace(',', ''))
            sales = float(str(fd_dict.get("sales", 1)).replace(',', ''))
            return f"{(op / sales * 100):.2f}"
        except Exception:
            return "N/A"

    ros_current = _pct(state.get("ros_current", _ros(fd)))
    ros_prior   = _pct(state.get("ros_prior",   _ros(fd_prior)))
    ros_prior2  = _pct(state.get("ros_prior2",  _ros(fd_prior2)))

    # ── org structure image ────────────────────────────────────────────────────
    org_image = None
    org_image_bytes = state.get("org_structure_image")
    if org_image_bytes:
        try:
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            tmp.write(org_image_bytes)
            tmp.close()
            org_image = InlineImage(tpl, tmp.name, width=Inches(6.0))
        except Exception:
            org_image = None

    # ── scope_of_report (static with company name injected) ──────────────────
    scope_of_report = (
        f"This Local File was prepared based on an understanding of Affiliated Party "
        f"transactions carried out by {company_short} in Fiscal Year {fiscal_year}. "
        f"In preparing this Local File, the information used includes:"
    )

    # ─────────────────────────────────────────────────────────────────────────
    context = {
        # Identity
        "company_name":        company_name,
        "company_short":       company_short,
        "company_name_no_pt":  company_name_no_pt,
        "fiscal_year":         fiscal_year,
        "brand_name":          brand_name,
        "company_address":     _safe(state.get("company_address")),
        "establishment_info":  _safe(state.get("establishment_info")),
        "employee_count":      _safe(state.get("employee_count"), "N/A"),
        "parent_company":      _safe(state.get("parent_company")),
        "parent_group":        _safe(state.get("parent_group")),
        "analysis_period":     _safe(state.get("analysis_period"), "2022-2024"),
        "scope_of_report":     scope_of_report,

        # Ownership & management
        "shareholders":        shareholders,
        "shareholders_source": _safe(state.get("shareholders_source")) or f"Source: Management {company_short}, 31 December {fiscal_year}",
        "management":          management,
        "management_source":   _safe(state.get("management_source")) or f"Source: Management {company_short}, 31 December {fiscal_year}",

        # Business
        "business_activities_description": _safe(state.get("business_activities_description")),
        "products":     products,
        "competitors":  competitors,
        "business_strategy":       _safe(state.get("business_strategy")),
        # Default: "not involved in any restructuring" when left blank
        "business_restructuring":  _safe(state.get("business_restructuring")) or (
            f"In Fiscal Year {fiscal_year}, {company_short} is not involved or affected by "
            f"any business restructuring or transfer of intangible assets within the Business "
            f"Group that is or has occurred in the previous year."
        ),
        "business_environment_overview":  _safe(state.get("business_environment_overview")),
        "business_environment_sources":   state.get("business_environment_sources") or [],

        # Transactions
        "affiliated_parties":          affiliated_parties,
        "affiliated_transactions":     affiliated_transactions,
        "independent_transactions":    independent_transactions,
        "transaction_details_text":    _strip_md(_safe(state.get("transaction_details_text"))),
        "pricing_policy":              _strip_md(_safe(state.get("pricing_policy"))),
        # Sec 4.1.2.4 — static template text
        "copy_of_agreement":           "A copy of the agreement/contract for the affiliated transaction is attached in Appendix 2.",

        # Section 4.1.2.1 (AI) — plain prose, no labels/bullets
        "background_transaction":           _strip_md(_safe(state.get("background_transaction"))),

        # Supply chain management (AI)
        "supply_chain_management":          _strip_md(_safe(state.get("supply_chain_management"))),

        # Functional analysis (AI) — strip markdown
        "functional_analysis_narrative":    _strip_md(_safe(state.get("functional_analysis_narrative"))),
        "business_characterization_text":   _strip_md(_safe(state.get("business_characterization_text"))),

        # Industry analysis (AI)
        "industry_analysis_global":      _strip_md(_safe(state.get("industry_analysis_global"))),
        "industry_global_sources":       state.get("industry_global_sources") or [],
        "industry_analysis_indonesia":   _strip_md(_safe(state.get("industry_analysis_indonesia"))),
        "industry_indonesia_sources":    state.get("industry_indonesia_sources") or [],
        # Dynamic industry heading — derived from products list (falls back to 'Industry')
        "industry_name": _safe(state.get("industry_name")) or (
            ", ".join(p.get("name", "") for p in state.get("products", []) if p.get("name"))
            or "Industry"
        ),
        # Company location analysis (AI) — Sec 5.1.2.1
        "company_location_analysis":              _strip_md(_safe(state.get("company_location_analysis"))),
        "efficiency_excellence_company_locations": _strip_md(_safe(state.get("company_location_analysis"))),
        "company_location_sources":               state.get("company_location_sources") or [],
        # Industry regulations (AI) — Sec 5.1.2.2
        "industry_regulations_text":    _strip_md(_safe(state.get("industry_regulations_text"))),
        "regulations_affecting_industry": _strip_md(_safe(state.get("industry_regulations_text"))),
        "industry_regulations_sources": state.get("industry_regulations_sources") or [],
        # Business environment (AI) — Sec 5.1.3
        "business_environment_overview": _strip_md(_safe(state.get("business_environment_overview"))),

        # TP method (AI)
        "method_selection_justification": _strip_md(_safe(state.get("method_selection_justification"))),
        "tested_party":                 _safe(state.get("tested_party"), company_short),
        "selected_method":              _safe(state.get("selected_method"), "TNMM"),
        "selected_pli":                 _safe(state.get("selected_pli"), "ROS"),
        "pli_selection_rationale":      _strip_md(_safe(state.get("pli_selection_rationale"))),

        # Comparability factors (Table 5.1)
        "comparability_factors": [
            {"factor": _safe(cf.get("factor")), "description": _safe(cf.get("description"))}
            for cf in state.get("comparability_factors", [])
            if cf.get("factor")
        ],

        # Comparable search
        "search_criteria_results": search_criteria_results,
        "rejection_matrix":        rejection_matrix,
        "comparable_companies":    comparable_companies,
        "comparability_analysis_narrative": _strip_md(_safe(state.get("comparability_analysis_narrative"))),

        # IQR
        "q1":                    q1,
        "median":                median,
        "q3":                    q3,
        "tested_party_ratio":    tested_party_ratio,
        "ros_current":           ros_current,
        "ros_prior":             ros_prior,
        "ros_prior2":            ros_prior2,

        # Financial data (raw dicts — for advanced template use)
        "financial_data":        fd,
        "financial_data_prior":  fd_prior,
        "financial_data_prior2": fd_prior2,

        # P/L table
        "pl_table": pl_table,

        # P/L overview (AI-generated bullet points)
        "pl_overview_text": _strip_md(_safe(state.get("pl_overview_text"))),
        "pl_bullets": [
            line.lstrip("-•").strip()
            for line in _strip_md(_safe(state.get("pl_overview_text"))).splitlines()
            if line.strip().startswith(("-", "•"))
        ],

        # Non-financial
        "non_financial_events": _strip_md(_safe(state.get("non_financial_events"))),

        # Org structure
        "org_structure_description": _safe(state.get("org_structure_description")),
        "org_structure_departments": state.get("org_structure_departments", []),
        "org_structure_image":       org_image,

        # Executive summary & conclusion (AI)
        "executive_summary": _strip_md(_safe(state.get("executive_summary"))),
        "conclusion_text":   _strip_md(_safe(state.get("conclusion_text"))),
    }

    return context


# ─────────────────────────────────────────────────────────────────────────────
# Main render function
# ─────────────────────────────────────────────────────────────────────────────

def render_tp_document(state: dict, output_path: str) -> None:
    """
    Render the Jinja2 DOCX template with data from `state`
    and write the result to `output_path`.
    """
    # Auto-generate template if missing (docs_template/ is in .gitignore)
    if not TEMPLATE_PATH.exists():
        import subprocess, sys
        script = _HERE / "apply_jinja_template.py"
        if not script.exists():
            raise FileNotFoundError(
                f"Neither template nor generator found.\n"
                f"Expected: {TEMPLATE_PATH}"
            )
        result = subprocess.run(
            [sys.executable, str(script)],
            capture_output=True, text=True, cwd=str(_HERE)
        )
        if not TEMPLATE_PATH.exists():
            raise FileNotFoundError(
                f"Template auto-generation failed.\n"
                f"STDOUT: {result.stdout}\nSTDERR: {result.stderr}"
            )

    tpl = DocxTemplate(str(TEMPLATE_PATH))
    context = build_context(state, tpl)
    tpl.render(context)
    tpl.save(output_path)  # Save first so the file is a valid DOCX

    # Post-process: rename headings → fill AI content → append source footnotes.
    from docx import Document as DocxDocument
    post_doc = DocxDocument(output_path)
    industry_name = context.get("industry_name", "Industry")
    _rename_headings(post_doc, industry_name)                      # step 1: rename old template headings
    _overwrite_section_bodies(post_doc, context)                   # step 2: fill AI/user content
    _add_industry_source_notes(post_doc, context)                  # step 3: append source footnotes
    _rebuild_appendix4(post_doc, context.get("comparable_companies", []))  # step 4: dynamic Appendix 4
    post_doc.save(output_path)

# ─────────────────────────────────────────────────────────────────────────────
# Post-render step 1: rename industry-specific template headings to generic ones
# ─────────────────────────────────────────────────────────────────────────────

# Old heading substrings to detect inside the template DOCX
_HEADING_OLD_KEYWORDS = [
    "global heavy equipment industry",
    "global industry analysis",   # in case template was already partially renamed
]
_HEADING_INDO_OLD_KEYWORDS = [
    "indonesian heavy equipment industry",
    "indonesian industry analysis",
]


def _rename_headings(doc, industry_name: str = "Industry") -> None:
    """
    Rename hardcoded template industry-analysis headings to names that reflect
    the company's actual industry (derived from its products list).

    The new heading text becomes:
      - Global:     "Global {industry_name} Industry Analysis"
      - Indonesian: "Indonesian {industry_name} Industry Analysis"

    Runs BEFORE _overwrite_section_bodies so the replacement keywords match.
    """
    import copy
    from docx.oxml.ns import qn

    # Cap industry_name length to keep headings readable
    _name = industry_name[:60].strip()
    new_global = f"Global {_name} Industry Analysis"
    new_indo   = f"Indonesian {_name} Industry Analysis"

    for para in doc.paragraphs:
        if not (para.style and para.style.name.startswith("Heading")):
            continue
        lower_text = para.text.lower().strip()

        if any(kw in lower_text for kw in _HEADING_OLD_KEYWORDS):
            target_text = new_global
        elif any(kw in lower_text for kw in _HEADING_INDO_OLD_KEYWORDS):
            target_text = new_indo
        else:
            continue

        # Clear runs, preserve rPr formatting from first run
        first_rPr = None
        if para.runs:
            rPr_el = para.runs[0]._r.find(qn("w:rPr"))
            if rPr_el is not None:
                first_rPr = copy.deepcopy(rPr_el)
        for child in list(para._p):
            if child.tag != qn("w:pPr"):
                para._p.remove(child)
        run = para.add_run(target_text)
        if first_rPr is not None:
            run._r.insert(0, first_rPr)


# ─────────────────────────────────────────────────────────────────────────────
# Post-render step 1b: append source footnotes after industry analysis sections
# ─────────────────────────────────────────────────────────────────────────────

def _add_industry_source_notes(doc, context: dict) -> None:
    """
    Insert footnote-style source references directly after the body text of the
    Global and Indonesian industry analysis sections.

    Visual style: thin top-border separator line + numbered URL entries in
    8pt Calibri italic — identical to _add_section_sources() in docx_export.py.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    from docx.text.paragraph import Paragraph

    global_sources = context.get("industry_global_sources", []) or []
    indo_sources   = context.get("industry_indonesia_sources", []) or []
    loc_sources    = context.get("company_location_sources", []) or []
    reg_sources    = context.get("industry_regulations_sources", []) or []
    env_sources    = context.get("business_environment_sources", []) or []

    # Map: heading keyword substring (lowercase) → (sources list, start_num)
    _n_glob = len(global_sources)
    _n_indo = len(indo_sources)
    _n_loc  = len(loc_sources)
    _n_reg  = len(reg_sources)
    _source_map = {
        "global":           (global_sources, 1),
        "indonesian":       (indo_sources,   _n_glob + 1),
        "efficiency":       (loc_sources,    _n_glob + _n_indo + 1),
        "regulation":       (reg_sources,    _n_glob + _n_indo + _n_loc + 1),
        "business environ": (env_sources,    _n_glob + _n_indo + _n_loc + _n_reg + 1),
    }

    para_list = list(doc.paragraphs)

    def _heading_level(para) -> int:
        name = para.style.name or ""
        if name.startswith("Heading"):
            try:
                return int(name.split()[-1])
            except ValueError:
                return 1
        return 0

    def _make_source_para(insert_after_el, text: str, is_separator: bool = False):
        """Create and insert a paragraph element after insert_after_el."""
        new_p = OxmlElement("w:p")
        insert_after_el.addnext(new_p)
        new_para = Paragraph(new_p, doc)

        pPr = new_para._p.get_or_add_pPr()
        if is_separator:
            pBdr = OxmlElement("w:pBdr")
            top  = OxmlElement("w:top")
            top.set(qn("w:val"),   "single")
            top.set(qn("w:sz"),    "6")
            top.set(qn("w:space"), "1")
            top.set(qn("w:color"), "000000")
            pBdr.append(top)
            pPr.append(pBdr)
            # spacing
            pPrSpacing = OxmlElement("w:spacing")
            pPrSpacing.set(qn("w:before"), "60")
            pPrSpacing.set(qn("w:after"),  "20")
            pPr.append(pPrSpacing)
        else:
            pPrSpacing = OxmlElement("w:spacing")
            pPrSpacing.set(qn("w:before"), "0")
            pPrSpacing.set(qn("w:after"),  "8")
            pPr.append(pPrSpacing)

        if text:
            run = new_para.add_run(text)
            run.font.name   = "Calibri"
            run.font.size   = Pt(8)
            run.font.italic = True

        return new_p

    for i, para in enumerate(para_list):
        lv = _heading_level(para)
        if lv == 0:
            continue
        lower = para.text.lower().strip()

        matched_key = None
        for kw in _source_map:
            if lower.startswith(kw):
                matched_key = kw
                break
        if matched_key is None:
            continue

        sources, start_num = _source_map[matched_key]
        if not sources:
            continue

        # Collect body paragraphs under this heading
        body_idxs = []
        for j in range(i + 1, len(para_list)):
            p2 = para_list[j]
            lv2 = _heading_level(p2)
            if lv2 and lv2 <= lv:
                break
            if p2.text.strip() and lv2 == 0:
                body_idxs.append(j)

        # Insert after the LAST body paragraph (or after heading if no body)
        anchor_el = para_list[body_idxs[-1]]._element if body_idxs else para._element

        # Insert in reverse order so each addnext pushes the previous one down
        # → first insert last entry, then entries before it, then separator
        inserts = []
        for idx, url in reversed(list(enumerate(sources, start_num))):
            inserts.append((f"[{idx}] {url}", False))
        inserts.append(("", True))  # separator (inserted last = appears first)

        current_anchor = anchor_el
        for text, is_sep in inserts:
            new_el = _make_source_para(current_anchor, text, is_separator=is_sep)
            current_anchor = new_el


# ─────────────────────────────────────────────────────────────────────────────
# Post-render step 2: overwrite static section body paragraphs with generated content
# ─────────────────────────────────────────────────────────────────────────────

# Mapping: heading keyword  →  context key that should replace its body text
_SECTION_REPLACEMENTS = [
    ("Background of Transaction to Affiliated Party", "background_transaction"),
    ("Business Activities and Operational Aspects",   "business_activities_description"),
    ("Business Strategy",                           "business_strategy"),
    # Use "of Transfer" to avoid matching the broader Heading 2 that also contains
    # "Business Restructuring" — that heading would cause the whole section 3.2 to
    # be overwritten with restructuring content (including Business Activities).
    ("Business Restructuring of Transfer",          "business_restructuring"),
    ("Pricing Policy",                              "pricing_policy"),
    ("Copy of Agreement",                           "copy_of_agreement"),
    ("Non-Financial Event",                         "non_financial_events"),
    # "Functional Analysis" intentionally removed — section uses static template text
    # with {{ company_short }} / {{ fiscal_year }} Jinja2 vars only; no AI override.
    # "Business Characterization" intentionally removed — user fills this manually via the app UI.
    # ("Business Characterization",                   "business_characterization_text"),
    ("Method Selection Justification",              "method_selection_justification"),
    ("PLI Selection Rationale",                     "pli_selection_rationale"),
    ("Comparability Analysis",                      "comparability_analysis_narrative"),
    # These keywords are matched AFTER _rename_headings runs.
    # The actual heading text will be "Global <industry_name> Industry Analysis"
    # — we match on the middle stable substring only.
    ("Global",      "industry_analysis_global"),
    ("Indonesian",  "industry_analysis_indonesia"),
    # Sec 5.1.2.1 — efficiency/excellence levels
    ("Efficiency and Excellence",   "company_location_analysis"),
    # Sec 5.1.2.2 — regulations affecting the industry
    ("Regulations Affecting",       "industry_regulations_text"),
    ("Overview of Profit/Loss",                     "pl_overview_text"),
]


def _set_para_text(para, text: str, reset_heading: bool = False) -> None:
    """
    Safely replace all run text in *para* with *text*.
    Preserves paragraph formatting (pPr) and the first run's character formatting.
    If reset_heading=True, changes Heading styles to 'Normal' so AI body text
    does not appear as numbered headings.
    """
    from docx.oxml.ns import qn

    # If this paragraph is a Heading style, reset it to Normal
    if reset_heading and para.style and para.style.name.startswith('Heading'):
        try:
            para.style = para.part.styles['Normal']
        except Exception:
            pass
        # Also clear pStyle in pPr so the inline style override is removed
        pPr = para._p.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                pPr.remove(pStyle)

    # Copy rPr from first run (if any) before clearing
    first_rPr = None
    if para.runs:
        first_r = para.runs[0]._r
        rPr_el = first_r.find(qn('w:rPr'))
        if rPr_el is not None:
            import copy
            first_rPr = copy.deepcopy(rPr_el)

    # Clear all existing runs (and field chars, bookmarks, etc.)
    # by removing every child that isn't w:pPr
    pPr = para._p.find(qn('w:pPr'))
    for child in list(para._p):
        if child.tag != qn('w:pPr'):
            para._p.remove(child)

    if not text:
        return

    # Add a fresh run with the text
    run = para.add_run(text)

    # When resetting from heading, don't carry over heading run formatting
    # (e.g. bold/caps from Heading rPr) — use plain run
    if not reset_heading and first_rPr is not None:
        run._r.insert(0, first_rPr)

    # Ensure xml:space=preserve so leading/trailing spaces are kept
    t_el = run._r.find(qn('w:t'))
    if t_el is not None:
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _rebuild_appendix4(doc, comparable_companies: list) -> None:
    """
    Replace hardcoded comparable company sections in Appendix 4 with one
    section per comparable company from the project state.

    Each section = numbered heading paragraph + cloned financial table
    (first existing table used as format template) + source paragraph.
    Company name and country are updated in the table header rows.
    Financial data rows are cleared (left blank) since per-company
    financial data is not stored in the project state.
    """
    from copy import deepcopy
    from lxml import etree
    from docx.oxml.ns import qn

    if not comparable_companies:
        return

    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    elements = list(body)

    # ── Find Appendix 4 ACTUAL heading (Heading1 style, not TOC entry) ─────────
    app4_elem = None
    for para in doc.paragraphs:
        if 'APPENDIX 4' not in para.text:
            continue
        style_el = para._element.find(f'.//{{{ns}}}pStyle')
        style_val = (style_el.get(f'{{{ns}}}val', '') if style_el is not None else '').lower()
        if 'heading' in style_val or style_val == '1':
            app4_elem = para._element
            # Don't break — take the LAST occurrence (body heading, not TOC)
    if app4_elem is None:
        return

    app4_idx = elements.index(app4_elem)

    # ── Collect all body elements belonging to Appendix 4 section ────────────
    section_elems = []
    fin_table_template = None

    for elem in elements[app4_idx + 1:]:
        tag = elem.tag.split('}')[-1]
        # Stop when we hit another Heading 1 (next appendix/section)
        if tag == 'p':
            style_el = elem.find(f'.//{{{ns}}}pStyle')
            style_val = (style_el.get(f'{{{ns}}}val', '') if style_el is not None else '').lower()
            text = ''.join(t.text or '' for t in elem.iter(f'{{{ns}}}t'))
            if ('heading' in style_val or style_val == '1') and text.strip() and 'APPENDIX 4' not in text:
                break
        section_elems.append(elem)
        if tag == 'tbl' and fin_table_template is None:
            fin_table_template = elem  # grab first table as format template

    if fin_table_template is None:
        return

    # ── Rescue any <w:sectPr> section-break properties before removing ─────────
    # Appendix 4 may live inside a landscape section. The <w:sectPr> embedded in
    # one of its paragraphs defines the page orientation for that section.
    # If we delete it, the landscape orientation bleeds into surrounding content.
    # Solution: find the sectPr, force portrait on it, and keep it attached to
    # a placeholder paragraph inserted right before the next heading.
    rescued_sect_pr = None
    for elem in section_elems:
        sp = elem.find(f'.//{{{ns}}}sectPr')
        if sp is not None:
            rescued_sect_pr = deepcopy(sp)
            # Force portrait orientation on the rescued sectPr
            pg_sz = rescued_sect_pr.find(f'{{{ns}}}pgSz')
            if pg_sz is not None:
                # Remove landscape orient attribute if present
                orient_attr = f'{{{ns}}}orient'
                if pg_sz.get(orient_attr) == 'landscape':
                    del pg_sz.attrib[orient_attr]
                # Swap w/h if w > h (landscape dimensions → portrait)
                w = pg_sz.get(f'{{{ns}}}w')
                h = pg_sz.get(f'{{{ns}}}h')
                if w and h and int(w) > int(h):
                    pg_sz.set(f'{{{ns}}}w', h)
                    pg_sz.set(f'{{{ns}}}h', w)
            break  # only need one

    # ── Remove all existing Appendix 4 content ───────────────────────────────
    for elem in section_elems:
        body.remove(elem)

    # ── Helper: clear financial data cells (keep structure, blank values) ─────
    HEADER_ROWS = 11  # rows 0–10 are header/metadata; row 11+ are financial data

    def _clone_table_for_company(cc: dict) -> etree._Element:
        clone = deepcopy(fin_table_template)
        rows = clone.findall(f'{{{ns}}}tr')
        name    = cc.get("name", "")
        country = cc.get("country", "")

        # Row 0 → company name
        if rows:
            for tc in rows[0].findall(f'.//{{{ns}}}tc'):
                for t in tc.findall(f'.//{{{ns}}}t'):
                    t.text = name

        # Row 1 → city / country
        if len(rows) > 1:
            for tc in rows[1].findall(f'.//{{{ns}}}tc'):
                for t in tc.findall(f'.//{{{ns}}}t'):
                    t.text = country

        # Rows 10+ (financial data) → blank values in non-label columns
        for row in rows[HEADER_ROWS:]:
            cells = row.findall(f'.//{{{ns}}}tc')
            for cell in cells[1:]:   # skip first column (account label)
                for t in cell.findall(f'.//{{{ns}}}t'):
                    t.text = ''

        return clone

    # ── Insert one section per comparable company ─────────────────────────────
    insert_ref = app4_elem   # keep track of where to insert next element

    for idx, cc in enumerate(comparable_companies):
        name = cc.get("name", f"Company {idx + 1}")

        # Numbered heading paragraph (List Paragraph style, e.g. "1. PT Astra...")
        name_para_xml = (
            f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:pPr><w:pStyle w:val="ListParagraph"/></w:pPr>'
            f'<w:r><w:rPr><w:b/></w:rPr>'
            f'<w:t xml:space="preserve">{idx + 1}.  {name}</w:t></w:r>'
            f'</w:p>'
        )
        name_elem = etree.fromstring(name_para_xml)

        # Financial table cloned from template
        table_elem = _clone_table_for_company(cc)

        # Source note paragraph
        source_para_xml = (
            f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:r><w:rPr><w:i/><w:sz w:val="18"/></w:rPr>'
            f'<w:t>Source: Bureau Van Dijk TP Catalyst Database</w:t></w:r>'
            f'</w:p>'
        )
        source_elem = etree.fromstring(source_para_xml)

        # Spacer paragraph
        spacer_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        spacer_elem = etree.fromstring(spacer_xml)

        # Insert: name → table → source → spacer
        for elem in [name_elem, table_elem, source_elem, spacer_elem]:
            insert_ref.addnext(elem)
            insert_ref = elem

    # ── Re-attach rescued sectPr in a closing paragraph ──────────────────────
    # This restores the section boundary so pages after Appendix 4 stay portrait.
    if rescued_sect_pr is not None:
        closing_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr/></w:p>'
        closing_p = etree.fromstring(closing_xml)
        closing_p.find(f'{{{ns}}}pPr').append(rescued_sect_pr)
        insert_ref.addnext(closing_p)


def _overwrite_section_bodies(doc, context: dict) -> None:
    """
    After docxtpl rendering, replace static body paragraphs under specific
    headings with the user/AI-generated content from *context*.

    Paragraph paragraphs are identified by appearing after a target heading
    and before the next heading of the same or higher level.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    paragraphs = doc.paragraphs   # top-level paragraphs only

    # Build index: para → position
    para_list = list(paragraphs)

    def _heading_level(para) -> int:
        """Return heading level 1-9 or 0 if not a heading."""
        name = para.style.name or ""
        if name.startswith("Heading"):
            try:
                return int(name.split()[-1])
            except ValueError:
                return 1
        return 0

    def _split_content(text: str) -> list[tuple[str, str]]:
        """
        Split AI-generated text into (content, style_hint) tuples.
        style_hint is 'bullet' for lines starting with '- ' or '• ',
        otherwise 'normal'.
        Uses \\a as primary paragraph separator, falls back to double-newline
        and single-newline splitting.
        """
        # First, try \a splitting (used by existing context vars)
        if '\a' in text:
            parts = [s.strip() for s in text.split('\a') if s.strip()]
        else:
            # Split on newlines but preserve logical paragraphs
            parts = []
            for line in text.split('\n'):
                line = line.strip()
                if line:
                    parts.append(line)

        result = []
        for part in parts:
            if part.startswith('- ') or part.startswith('• '):
                result.append((part[2:], 'bullet'))
            else:
                result.append((part, 'normal'))
        return result

    for heading_kw, ctx_key in _SECTION_REPLACEMENTS:
        new_text = (context.get(ctx_key) or "").strip()
        if not new_text:
            continue  # nothing to write — leave template content as-is

        # Find heading paragraph
        h_idx = None
        h_level = 0
        for i, p in enumerate(para_list):
            if _heading_level(p) and heading_kw.lower() in p.text.lower():
                h_idx = i
                h_level = _heading_level(p)
                break

        if h_idx is None:
            continue  # heading not found in this doc

        # Collect body paragraphs until next heading of same/higher level.
        body_idxs = []
        for i in range(h_idx + 1, len(para_list)):
            p = para_list[i]
            lv = _heading_level(p)
            if lv and lv <= h_level:
                break  # next section at same/higher level — stop
            if p.text.strip() and lv == 0:
                body_idxs.append(i)

        if not body_idxs:
            continue

        # Parse the AI content into sub-paragraphs
        sub_paras = _split_content(new_text)
        if not sub_paras:
            sub_paras = [(new_text, 'normal')]

        # Overwrite existing body paragraphs
        for j, idx in enumerate(body_idxs):
            if j < len(sub_paras):
                content, style_hint = sub_paras[j]
                _set_para_text(para_list[idx], content, reset_heading=True)
                # Apply bullet style if needed
                if style_hint == 'bullet':
                    try:
                        para_list[idx].style = doc.styles['List Paragraph']
                    except KeyError:
                        pass  # style not available in template
            else:
                _set_para_text(para_list[idx], "", reset_heading=True)  # clear extras

        # If AI generated MORE sub-paragraphs than body slots, insert new ones
        if len(sub_paras) > len(body_idxs):
            # Insert after the last body paragraph
            last_body_el = para_list[body_idxs[-1]]._element
            insert_after = last_body_el

            for extra_content, extra_style in sub_paras[len(body_idxs):]:
                # Create a new paragraph element
                new_p = OxmlElement('w:p')

                # Insert new_p right after insert_after in the XML tree
                insert_after.addnext(new_p)

                # Create a proper docx Paragraph wrapper
                from docx.text.paragraph import Paragraph
                new_para = Paragraph(new_p, doc)

                # Set style
                if extra_style == 'bullet':
                    try:
                        new_para.style = doc.styles['List Paragraph']
                    except KeyError:
                        new_para.style = doc.styles['Normal']
                else:
                    new_para.style = doc.styles['Normal']

                # Add text
                run = new_para.add_run(extra_content)

                insert_after = new_p

