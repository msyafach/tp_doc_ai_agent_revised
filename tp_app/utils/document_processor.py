"""
document_processor.py
======================
2-Tier retrieval strategy for TP document extraction:

  Tier 1 — PageIndex  (<  PAGE_THRESHOLD pages)
      Builds a hierarchical TOC tree from the PDF using the open-source
      PageIndex library (cloned at ./PageIndex).  Retrieval is done via
      LLM reasoning over the tree nodes — no embeddings needed.

  Tier 2 — Vector RAG (>= PAGE_THRESHOLD pages)
      Standard FAISS + RecursiveCharacterTextSplitter approach.
      Embedding model is HuggingFace all-MiniLM-L6-v2 (local, free)
      or OpenAI text-embedding-3-small if OPENAI_API_KEY is set.

Supported upload formats: PDF, DOCX, XLSX, TXT
Maximum file size enforced by the caller (app.py): 20 MB
"""

from __future__ import annotations

import os
import sys
import io
import json
import tempfile
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

# ── Add the cloned PageIndex repo to sys.path ──────────────────────────────────
_PI_ROOT = Path(__file__).resolve().parent.parent / "PageIndex"
if str(_PI_ROOT) not in sys.path:
    sys.path.insert(0, str(_PI_ROOT))

# ── Thresholds ─────────────────────────────────────────────────────────────────
PAGE_THRESHOLD = 20          # Use PageIndex if number of pages < this value

# ── Supported non-PDF types (loaded as plain text) ────────────────────────────
NON_PDF_EXTENSIONS = {".docx", ".xlsx", ".txt"}


# ══════════════════════════════════════════════════════════════════════════════
# Data-class holding the retrieval context passed to the extraction agent
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class RetrievalContext:
    strategy: str                          # "page_index" | "vector_rag"
    page_count: int = 0
    # PageIndex tier
    page_index_tree: Optional[dict] = None
    page_list: Optional[list] = None       # [(page_text, token_len), ...]
    # Vector RAG tier
    vectorstore: Any = None
    # Metadata
    filenames: list = field(default_factory=list)
    errors: list = field(default_factory=list)


# ══════════════════════════════════════════════════════════════════════════════
# PDF page count helper (PyPDF2, bundled with PageIndex deps)
# ══════════════════════════════════════════════════════════════════════════════

def _count_pdf_pages(file_bytes: bytes) -> int:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        return len(reader.pages)
    except Exception:
        return 0


# ══════════════════════════════════════════════════════════════════════════════
# Non-PDF text extraction helpers
# ══════════════════════════════════════════════════════════════════════════════

def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        raise RuntimeError(f"Cannot read DOCX: {e}")


def _extract_xlsx(file_bytes: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"--- Sheet: {ws.title} ---")
            for row in ws.iter_rows(values_only=True):
                line = "\t".join(str(c) if c is not None else "" for c in row)
                if line.strip():
                    parts.append(line)
        return "\n".join(parts)
    except Exception as e:
        raise RuntimeError(f"Cannot read XLSX: {e}")


def _extract_txt(file_bytes: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


# ══════════════════════════════════════════════════════════════════════════════
# Tier 1 — PageIndex  (short PDF)
# ══════════════════════════════════════════════════════════════════════════════

def _build_page_index(pdf_bytes: bytes, openai_api_key: str) -> tuple[dict, list]:
    """
    Builds a PageIndex tree for a short PDF (<PAGE_THRESHOLD pages).

    Returns (tree_dict, page_list) where page_list = [(text, tokens), ...]
    The tree_dict has shape: {"doc_name": ..., "structure": [...]}
    """
    from pageindex.page_index import page_index_main
    from pageindex.utils import ConfigLoader, get_page_tokens

    # PageIndex reads CHATGPT_API_KEY from env — bridge our key
    os.environ["CHATGPT_API_KEY"] = openai_api_key

    loader = ConfigLoader()
    opt = loader.load({
        "model": "gpt-4o-mini",
        "toc_check_page_num": 20,
        "max_page_num_each_node": 10,
        "max_token_num_each_node": 20000,
        "if_add_node_id": "yes",
        "if_add_node_summary": "yes",
        "if_add_doc_description": "no",
        "if_add_node_text": "yes",   # include text so we can pass it to LLM
    })

    pdf_stream = io.BytesIO(pdf_bytes)
    tree = page_index_main(pdf_stream, opt)
    page_list = get_page_tokens(io.BytesIO(pdf_bytes), model=opt.model)
    return tree, page_list


# ══════════════════════════════════════════════════════════════════════════════
# Tier 2 — Vector RAG  (long documents, or non-PDF files)
# ══════════════════════════════════════════════════════════════════════════════

def _get_embeddings():
    """Returns the best available embeddings — OpenAI first, then HuggingFace."""
    if os.environ.get("OPENAI_API_KEY"):
        from langchain_openai import OpenAIEmbeddings
        return OpenAIEmbeddings(
            model="text-embedding-3-small",
            api_key=os.environ["OPENAI_API_KEY"],
        )
    from langchain_huggingface import HuggingFaceEmbeddings
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")


def _build_vectorstore(text: str):
    """Chunk text and embed into an in-memory FAISS store."""
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import FAISS
    from langchain.schema import Document

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    chunks = splitter.split_text(text)
    if not chunks:
        raise ValueError("No text content could be extracted from the document.")

    docs = [Document(page_content=c) for c in chunks]
    embeddings = _get_embeddings()
    return FAISS.from_documents(docs, embeddings)


# ══════════════════════════════════════════════════════════════════════════════
# Main pipeline — called from app.py
# ══════════════════════════════════════════════════════════════════════════════

def process_uploaded_files(uploaded_files) -> tuple[Optional[RetrievalContext], list[str]]:
    """
    Process one or more Streamlit UploadedFile objects.

    Returns (RetrievalContext, error_list).
    RetrievalContext is None if no usable content was found.
    """
    errors: list[str] = []
    all_text_parts: list[str] = []      # for Vector RAG
    pdf_files: list[tuple[str, bytes]] = []  # (name, bytes) — PDF-only for PageIndex
    filenames: list[str] = []

    openai_key = os.environ.get("OPENAI_API_KEY", "")

    for uf in uploaded_files:
        ext = Path(uf.name).suffix.lower()
        raw = uf.read()
        filenames.append(uf.name)

        try:
            if ext == ".pdf":
                pdf_files.append((uf.name, raw))
                # Also extract text for the vector fallback path
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(raw))
                text = "\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
                all_text_parts.append(f"[FILE: {uf.name}]\n{text}")

            elif ext == ".docx":
                text = _extract_docx(raw)
                all_text_parts.append(f"[FILE: {uf.name}]\n{text}")

            elif ext == ".xlsx":
                text = _extract_xlsx(raw)
                all_text_parts.append(f"[FILE: {uf.name}]\n{text}")

            elif ext == ".txt":
                text = _extract_txt(raw)
                all_text_parts.append(f"[FILE: {uf.name}]\n{text}")

            else:
                errors.append(f"{uf.name}: unsupported extension '{ext}'")

        except Exception as exc:
            errors.append(f"{uf.name}: {exc}")

    if not all_text_parts and not pdf_files:
        return None, errors

    # ── Decide strategy ───────────────────────────────────────────────────────
    # PageIndex is only eligible if:
    #   1. There is exactly 1 PDF upload (multi-file → merge → vector RAG)
    #   2. That PDF has < PAGE_THRESHOLD pages
    #   3. An OpenAI key is available (PageIndex needs GPT)
    use_page_index = False
    page_count = 0

    if (
        len(pdf_files) == 1
        and len(uploaded_files) == 1          # single PDF only
        and openai_key
    ):
        page_count = _count_pdf_pages(pdf_files[0][1])
        if 0 < page_count < PAGE_THRESHOLD:
            use_page_index = True

    # ── Tier 1: PageIndex ─────────────────────────────────────────────────────
    if use_page_index:
        try:
            fname, raw = pdf_files[0]
            tree, page_list = _build_page_index(raw, openai_key)
            return RetrievalContext(
                strategy="page_index",
                page_count=page_count,
                page_index_tree=tree,
                page_list=page_list,
                filenames=filenames,
                errors=errors,
            ), errors
        except Exception as exc:
            errors.append(
                f"PageIndex failed ({exc}), falling back to Vector RAG."
            )
            # fall through to vector RAG

    # ── Tier 2: Vector RAG ────────────────────────────────────────────────────
    merged_text = "\n\n".join(all_text_parts)
    if not merged_text.strip():
        return None, errors

    try:
        vs = _build_vectorstore(merged_text)
    except Exception as exc:
        errors.append(f"Embedding failed: {exc}")
        return None, errors

    return RetrievalContext(
        strategy="vector_rag",
        page_count=page_count,
        vectorstore=vs,
        filenames=filenames,
        errors=errors,
    ), errors
