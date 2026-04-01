"""
fix_header_border.py
====================
Restores the blue separator line (w:drawing shape) that was removed
when fixing the header alignment. The drawing is extracted from the
backup file and re-inserted into the current template's headers.
"""

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy

BAK = "TP_Local_File_TEMPLATE.docx.bak"
CUR = "TP_Local_File_TEMPLATE.docx"

# ── Step 1: Extract drawing elements from backup headers ──────────────────────

doc_bak = Document(BAK)

# Collect unique drawing XMLs from backup headers (by section)
# We'll use section 1 as the reference since it's the main repeating header
drawings_by_section = {}

for i, s in enumerate(doc_bak.sections):
    hdr = s.header
    for p in hdr.paragraphs:
        xml = etree.tostring(p._p).decode()
        if "w:drawing" in xml:
            # Find w:r children that contain the drawing
            drawing_runs = []
            for child in p._p:
                child_xml = etree.tostring(child).decode()
                if "w:drawing" in child_xml:
                    drawing_runs.append(copy.deepcopy(child))
            if drawing_runs:
                drawings_by_section[i] = drawing_runs
                print(f"  Sec{i}: found {len(drawing_runs)} drawing run(s)")
            break

if not drawings_by_section:
    print("ERROR: No drawings found in backup!")
    exit(1)

# Use section 1 drawing as the reference (or first available)
ref_sec = min(drawings_by_section.keys())
ref_drawings = drawings_by_section[ref_sec]
print(f"Using section {ref_sec} drawing as reference.")

# ── Step 2: Insert drawings into current template headers ──────────────────────

doc_cur = Document(CUR)

fixed_count = 0
for i, s in enumerate(doc_cur.sections):
    hdr = s.header
    for p in hdr.paragraphs:
        full = "".join(r.text for r in p.runs)
        if full.strip():  # main header paragraph
            # Check if it already has a drawing
            xml = etree.tostring(p._p).decode()
            if "w:drawing" in xml:
                print(f"  Sec{i}: already has drawing, skipping.")
                continue

            # Find the best drawing source for this section
            # Use section-specific if available, else use reference
            src_drawings = drawings_by_section.get(i, ref_drawings)

            # Append drawing run(s) at end of paragraph (after existing runs)
            for dr in src_drawings:
                dr_copy = copy.deepcopy(dr)
                p._p.append(dr_copy)

            fixed_count += 1
            print(f"  Sec{i}: drawing inserted OK.")
            break

print(f"\nFixed {fixed_count} sections.")
doc_cur.save(CUR)
print(f"Saved to {CUR}")
