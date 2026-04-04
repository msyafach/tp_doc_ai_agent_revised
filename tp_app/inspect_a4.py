# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx", "lxml"]
# ///
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
from docx import Document
from docx.oxml.ns import qn

doc = Document("TP_Local_File_TEMPLATE.docx")
print(f"Total tables: {len(doc.tables)}")

# Inspect tables 19-25 (Appendix 3 and 4 area)
for idx in range(19, 26):
    table = doc.tables[idx]
    tbl = table._tbl
    nrows = len(table.rows)
    ncols = len(table.columns)
    # Get first 5 rows
    print(f"\n=== Table {idx} ({nrows}r x {ncols}c) ===")
    for r, row in enumerate(table.rows[:8]):
        tcs = row._tr.findall(qn("w:tc"))
        cells = []
        for tc in tcs:
            text = "".join(
                t.text or "" for p in tc.findall(qn("w:p")) for r2 in p.findall(qn("w:r")) for t in r2.findall(qn("w:t"))
            )
            cells.append(text.strip()[:25])
        print(f"  row{r}: {cells}")
    if nrows > 8:
        print(f"  ... ({nrows - 8} more rows)")
