"""
apply_jinja_template.py
========================
Copies the original PT SP DOCX and applies Jinja2 placeholders
throughout so the file becomes a docxtpl-compatible template.

Run:  uv run python apply_jinja_template.py
Output: docs_template/TP_Local_File_jinja2.docx
"""

import re
import shutil
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SRC = r"docs_template/PT Sany Perkasa_Local File FY 2024_AI TPDoc_220126_TEMPLATE (2).docx"
DST = r"docs_template/TP_Local_File_jinja2.docx"

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def para_full_text(para):
    return "".join(r.text for r in para.runs)


def replace_in_para(para, old, new):
    """Replace `old` with `new` in a paragraph, handling split runs."""
    full = para_full_text(para)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    # Put everything in the first run, clear the rest
    runs = para.runs
    if not runs:
        return False
    runs[0].text = new_full
    for r in runs[1:]:
        r.text = ""
    return True


def replace_in_cell(cell, old, new):
    """Replace text in all paragraphs of a cell."""
    changed = False
    for para in cell.paragraphs:
        if replace_in_para(para, old, new):
            changed = True
    return changed


def replace_all(body_paras, tables, pairs):
    """
    Apply a list of (old, new) pairs to every paragraph and table cell.
    Processes in order — later pairs can reference results of earlier ones.
    """
    for old, new in pairs:
        for para in body_paras:
            replace_in_para(para, old, new)
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_cell(cell, old, new)


def set_cell_text(cell, text):
    """Overwrite all text in a cell's first paragraph (keeps formatting)."""
    for i, para in enumerate(cell.paragraphs):
        runs = para.runs
        if runs:
            if i == 0:
                runs[0].text = text
                for r in runs[1:]:
                    r.text = ""
            else:
                for r in runs:
                    r.text = ""
        else:
            if i == 0:
                para.add_run(text)


def convert_table_to_loop(table, loop_var, fields,
                           header_rows=1, skip_last_n_rows=0,
                           for_tag=None, endfor_tag=None):
    """
    Turn a data-carrying table into a docxtpl row-loop template.
    - header_rows: number of rows at the top to keep as-is (caption + column headings)
    - Keeps the FIRST data row as the template row, deletes the rest
    - Replaces cell content of the template row with {{ loop_var.field }}
    - Wraps the template row with {%tr for %} / {%tr endfor %} tags

    `fields` is a list aligned to table.columns, e.g.:
        ['loop.index', 'item.name', 'item.country']
    Use None to leave a cell unchanged.
    """
    total_rows = len(table.rows)
    data_start = header_rows
    data_end = total_rows - skip_last_n_rows

    if data_end <= data_start:
        return   # nothing to template

    # Delete extra data rows (keep only the first data row as template)
    rows_to_delete = table.rows[data_start + 1 : data_end]
    tbl_elem = table._tbl
    for row in rows_to_delete:
        tbl_elem.remove(row._tr)

    # Now set template row cells
    template_row = table.rows[data_start]
    for col_idx, field in enumerate(fields):
        if field is None:
            continue
        if col_idx < len(template_row.cells):
            cell = template_row.cells[col_idx]
            set_cell_text(cell, '{{ ' + field + ' }}')

    # Inject {%tr for %} before template row and {%tr endfor %} after
    for_instr = for_tag or f'{{% tr for {loop_var} in {loop_var}s %}}'
    end_instr = endfor_tag or '{%tr endfor %}'

    tr = template_row._tr
    # Build two w:tr elements carrying just a single w:tc with the tag text
    def make_tag_row(text):
        new_tr = etree.SubElement(tbl_elem, qn('w:tr'))
        new_tc = etree.SubElement(new_tr, qn('w:tc'))
        new_p  = etree.SubElement(new_tc, qn('w:p'))
        new_r  = etree.SubElement(new_p,  qn('w:r'))
        new_t  = etree.SubElement(new_r,  qn('w:t'))
        new_t.text = text
        # Must appear in correct position – move to right place after
        return new_tr

    for_row = make_tag_row(for_instr)
    end_row = make_tag_row(end_instr)

    # Reorder: for_row → template_row → end_row
    tbl_elem.remove(for_row)
    tbl_elem.remove(end_row)
    tbl_elem.insert(list(tbl_elem).index(tr), for_row)
    tbl_elem.insert(list(tbl_elem).index(tr) + 1, end_row)


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

print(f"Copying {SRC} -> {DST}")
shutil.copy2(SRC, DST)

doc = Document(DST)
paras  = doc.paragraphs
tables = doc.tables

# ═════════════════════════════════════════════════════════════════════════════
# 1. GLOBAL TEXT REPLACEMENTS (order matters — most specific first)
# ═════════════════════════════════════════════════════════════════════════════
pairs = [
    # --- company / year tokens ---
    ("PT Sany Perkasa",       "{{ company_name }}"),
    ("PT SP",                 "{{ company_short }}"),
    ("FY 2024",               "FY {{ fiscal_year }}"),
    ("FY 2023",               "FY {{ fiscal_year|int - 1 }}"),
    ("FY 2022",               "FY {{ fiscal_year|int - 2 }}"),
    ("31 December 2024",      "31 December {{ fiscal_year }}"),
    ("31 December 2023",      "31 December {{ fiscal_year|int - 1 }}"),
    ("Fiscal Year 2024",      "Fiscal Year {{ fiscal_year }}"),
    ("fiscal year 2024",      "fiscal year {{ fiscal_year }}"),
    ("in 2024",               "in {{ fiscal_year }}"),
    ("In 2024",               "In {{ fiscal_year }}"),
    ("in FY 2024",            "in FY {{ fiscal_year }}"),
    ("In FY 2024",            "In FY {{ fiscal_year }}"),
    ("2024 and 2023",         "{{ fiscal_year }} and {{ fiscal_year|int - 1 }}"),
    ("2022-2024",             "{{ analysis_period }}"),
    ("2020-2022",             "{{ analysis_period }}"),
    # --- IQR / ratio numbers ---
    ("1st quartile 2.14% to 3rd quartile 4.05%, with a median of 3.13%",
     "1st quartile {{ q1 }}% to 3rd quartile {{ q3 }}%, with a median of {{ median }}%"),
    ("2.14%",   "{{ q1 }}%"),
    ("4.05%",   "{{ q3 }}%"),
    ("3.13%",   "{{ median }}%"),
    ("2.48%",   "{{ tested_party_ratio }}%"),
    # --- parent/group names ---
    ("Sany Group",            "{{ parent_group }}"),
    ("Sany Southeast Asia Pte, Ltd",  "{{ shareholders[0].name }}"),
    # --- headcount ---
    ("1,034",                 "{{ employee_count }}"),
]

print("Applying global text replacements...")
replace_all(paras, tables, pairs)

# ═════════════════════════════════════════════════════════════════════════════
# 2. LONG NARRATIVE PARAGRAPH REPLACEMENTS
#    (replace the content of specific paragraphs that are AI-generated)
# ═════════════════════════════════════════════════════════════════════════════

def replace_para_starting_with(paras, starts_with, replacement, partial=False):
    """Replace the first paragraph whose text starts with `starts_with`."""
    for para in paras:
        full = para_full_text(para)
        if partial:
            if starts_with.lower() in full.lower():
                runs = para.runs
                if runs:
                    runs[0].text = replacement
                    for r in runs[1:]:
                        r.text = ""
                return True
        else:
            if full.strip().startswith(starts_with):
                runs = para.runs
                if runs:
                    runs[0].text = replacement
                    for r in runs[1:]:
                        r.text = ""
                return True
    return False

narrative_replacements = [
    # Executive summary Chapter 1
    ("PT SP is a subsidiary of Sany Group",
     "{{ executive_summary }}"),
    # Scope of report
    ("This Local File was prepared based on an understanding",
     "{{ scope_of_report }}"),
    # Business activities
    ('PT Sany Perkasa ("PT SP") was established in Jakarta based on Deed No. 1419',
     "{{ establishment_info }}"),
    ("PT SP is engaged in the business of trading heavy equipment",
     "{{ business_activities_description }}"),
    # Business strategy
    ("For more than three decades",
     "{{ business_strategy }}"),
    # Business restructuring
    ("In the Fiscal Year 2024, PT SP is not involved or affected by any business restructuring",
     "{{ business_restructuring }}"),
    # Industry - global
    ("Heavy equipment refers to large, powerful machines that are specifically designed",
     "{{ industry_analysis_global }}"),
    ("The global construction heavy equipment market size was valued at USD 208.86 billion in 2024 and is estimated",
     ""),   # second para of global — will be merged into above
    # Industry - Indonesia
    ("The heavy equipment industry in Indonesia has a very important role",
     "{{ industry_analysis_indonesia }}"),
    # Functional analysis narrative
    ("Based on chapter 4.1 of this Local File, the party tested",
     "{{ functional_analysis_narrative }}"),
    # Business characterization
    ("Based on the functional analysis explained above, it can be concluded",
     "{{ business_characterization_text }}"),
    # Pricing policy
    ("The purchase price is determined between related parties",
     "{{ pricing_policy }}"),
    # Transaction detail narrative
    ("In 2024, PT SP purchased of goods from affiliated parties",
     "{{ transaction_details_text }}"),
    # Method selection
    ("Based on the transfer pricing method recognized by the OECD",
     "{{ method_selection_justification }}"),
    # PLI selection rationale
    ("Based on PT SP's functional analysis in relation to related party transactions, the most appropriate PLI",
     "{{ pli_selection_rationale }}"),
    # Comparability analysis
    ("Based on table 5.6, the arm's length range of ROS ratios",
     "{{ comparability_analysis_narrative }}"),
    # Conclusion Chapter 5.8
    ("In conducting transfer pricing analysis, TNMM was selected as the most appropriate method",
     "{{ conclusion_text }}"),
    # Non-financial events
    ("In FY 2024, as part of its marketing activities",
     "{{ non_financial_events }}"),
]

print("Applying narrative paragraph replacements...")
for starts, repl in narrative_replacements:
    replace_para_starting_with(paras, starts, repl)


# ═════════════════════════════════════════════════════════════════════════════
# 3. TABLE CAPTION REPLACEMENTS
# ═════════════════════════════════════════════════════════════════════════════
table_caption_pairs = [
    ("Table 3.1 Ownership Structure PT SP in FY 2024",
     "Table {{ loop.index }} Ownership Structure {{ company_short }} in FY {{ fiscal_year }}"),
    ("Table 3.2 Management PT SP in FY 2024",
     "Table {{ loop.index }} Management {{ company_short }} in FY {{ fiscal_year }}"),
    ("Table 3.3 Details of Products PT SP",
     "Table {{ loop.index }} Details of Products {{ company_short }}"),
    ("Table 3.4 List of Competitors",
     "Table {{ loop.index }} List of Competitors"),
    ("Table 4.1 List of PT SP's Affiliated Transactions in FY 2024",
     "Table {{ loop.index }} List of {{ company_short }}'s Affiliated Transactions in FY {{ fiscal_year }}"),
    ("Table 4.2 List of PT SP's Independent Transactions in FY 2024",
     "Table {{ loop.index }} List of {{ company_short }}'s Independent Transactions in FY {{ fiscal_year }}"),
    ("Table 5.3 Summary of Comparable Company Search Criteria",
     "Table {{ loop.index }} Summary of Comparable Company Search Criteria"),
    ("Table 5.4 Manual Review for Comprable Companies",
     "Table {{ loop.index }} Manual Review for Comparable Companies"),
    ("Table 5.5 Selected Comparable Companies",
     "Table {{ loop.index }} Selected Comparable Companies"),
    ("Table 5.6 Inter-quartile Range of the ROS Ratio of Comparable Companies",
     "Table {{ loop.index }} Inter-quartile Range of the ROS Ratio of Comparable Companies"),
    ("Table 5.7 Calculation of PT SP's ROS Ratio in FY 2022-2024",
     "Table {{ loop.index }} Calculation of {{ company_short }}'s ROS Ratio in FY {{ analysis_period }}"),
    ("Table 6.1 Profit/Loss Summary of PT SP for Fiscal Year 2024 and Fiscal Year 2023",
     "Table {{ loop.index }} Profit/Loss Summary of {{ company_short }} for Fiscal Year {{ fiscal_year }} and {{ fiscal_year|int - 1 }}"),
    ("Table 4.3 PT SP Functional Analysis",
     "Table {{ loop.index }} {{ company_short }} Functional Analysis"),
    ("Table 5.1 Comparability Analysis of PT SP Affiliated Transactions FY 2024",
     "Table {{ loop.index }} Comparability Analysis of {{ company_short }} Affiliated Transactions FY {{ fiscal_year }}"),
    ("Table 5.2  Comparability Analysis of PT SP Purchase Transactions",
     "Table {{ loop.index }} Comparability Analysis of {{ company_short }} Purchase Transactions"),
    # Source lines
    ("Source: Management PT SP, 31 December 2024",
     "Source: Management {{ company_short }}, 31 December {{ fiscal_year }}"),
    ("Source: Management PT SP, 31 December 2024 and 2023",
     "Source: Management {{ company_short }}, 31 December {{ fiscal_year }} and {{ fiscal_year|int - 1 }}"),
    ("Source: Management PT SP FY 2024",
     "Source: Management {{ company_short }} FY {{ fiscal_year }}"),
    ("Source: Website PT SP",
     "Source: Website {{ company_short }}"),
    ("Source: Management PT SP, 31 December 2024 and 2023",
     "Source: Management {{ company_short }}, {{ fiscal_year }} and {{ fiscal_year|int - 1 }}"),
]

print("Applying table caption replacements...")
replace_all(paras, tables, table_caption_pairs)


# ═════════════════════════════════════════════════════════════════════════════
# 4. CONVERT DATA TABLES TO JINJA2 LOOPS
#    Table index reference (from _tables.txt scan):
#    T02 = Shareholders (4 cols: name, shares, capital, pct)
#    T03 = Management (2 cols: position, name) — grouped sections
#    T04 = Products (2 cols: name, description)
#    T05 = Competitors (3 cols: no, name, description)
#    T06 = Affiliated transactions (7 cols: no, party, jurisdiction, affiliation, type, value, note)
#    T07 = Independent transactions (5 cols: no, party, location, type, value)
#    T13 = Search criteria (3 cols: no, criteria, result)
#    T15 = Selected comparables (3 cols: no, name, activities)
#    T18 = P/L summary (4 cols: account, FY current, FY prior, performance)
#    T19 = Affiliated Parties list in Appendix 2 (4 cols)
#    T20 = Comparable Companies in Appendix 3 (3 cols)
# ═════════════════════════════════════════════════════════════════════════════

print("Converting table T02 (Shareholders) to Jinja loop...")
# T02: caption row [0] + header row [1] → keep. Data rows [2..4] → loop
try:
    convert_table_to_loop(
        tables[2],
        loop_var='sh',
        fields=['sh.name', 'sh.shares', 'sh.capital', 'sh.percentage'],
        header_rows=2,
        for_tag='{% tr for sh in shareholders %}',
        endfor_tag='{% tr endfor %}',
    )
except Exception as e:
    print(f"  T02 warning: {e}")

print("Converting table T04 (Products) to Jinja loop...")
# T04: caption row [0] + header row [1] → keep. Data rows [2..4] → loop
try:
    convert_table_to_loop(
        tables[4],
        loop_var='prod',
        fields=['prod.name', 'prod.description'],
        header_rows=2,
        for_tag='{% tr for prod in products %}',
        endfor_tag='{% tr endfor %}',
    )
except Exception as e:
    print(f"  T04 warning: {e}")

print("Converting table T05 (Competitors) to Jinja loop...")
try:
    convert_table_to_loop(
        tables[5],
        loop_var='comp',
        fields=['loop.index', 'comp.name', 'comp.description'],
        header_rows=2,
        for_tag='{% tr for comp in competitors %}',
        endfor_tag='{% tr endfor %}',
    )
except Exception as e:
    print(f"  T05 warning: {e}")

print("Converting table T06 (Affiliated Transactions) to Jinja loop...")
try:
    convert_table_to_loop(
        tables[6],
        loop_var='aff',
        # cols: no, party, jurisdiction, affiliation type, transaction type, value, note
        fields=['loop.index', 'aff.name', 'aff.country', 'aff.affiliation_type',
                'aff.transaction_type', 'aff.value', 'aff.note'],
        header_rows=2,
        for_tag='{% tr for aff in affiliated_transactions %}',
        endfor_tag='{% tr endfor %}',
    )
except Exception as e:
    print(f"  T06 warning: {e}")

print("Converting table T07 (Independent Transactions) to Jinja loop...")
try:
    convert_table_to_loop(
        tables[7],
        loop_var='ind',
        fields=['loop.index', 'ind.name', 'ind.country', 'ind.transaction_type', 'ind.value'],
        header_rows=2,
        for_tag='{% tr for ind in independent_transactions %}',
        endfor_tag='{% tr endfor %}',
    )
except Exception as e:
    print(f"  T07 warning: {e}")

print("Converting table T13 (Search Criteria) to Jinja loop...")
try:
    convert_table_to_loop(
        tables[13],
        loop_var='sc',
        fields=['sc.step', 'sc.criteria', 'sc.result_count'],
        header_rows=2,
        for_tag='{% tr for sc in search_criteria_results %}',
        endfor_tag='{% tr endfor %}',
    )
except Exception as e:
    print(f"  T13 warning: {e}")

print("Converting table T15 (Selected Comparables) to Jinja loop...")
try:
    convert_table_to_loop(
        tables[15],
        loop_var='c',
        fields=['loop.index', 'c.name', 'c.description'],
        header_rows=2,
        for_tag='{% tr for c in comparable_companies %}',
        endfor_tag='{% tr endfor %}',
    )
except Exception as e:
    print(f"  T15 warning: {e}")

print("Converting table T18 (P/L Summary) to Jinja loop...")
try:
    convert_table_to_loop(
        tables[18],
        loop_var='pl',
        fields=['pl.account', 'pl.fy_current', 'pl.fy_prior', 'pl.performance'],
        header_rows=2,
        for_tag='{% tr for pl in pl_table %}',
        endfor_tag='{% tr endfor %}',
    )
except Exception as e:
    print(f"  T18 warning: {e}")

print("Converting table T19 (Appendix 2 — Affiliated Parties) to Jinja loop...")
try:
    convert_table_to_loop(
        tables[19],
        loop_var='ap',
        fields=['loop.index', 'ap.name', 'ap.country', 'ap.affiliation_type',
                'ap.transaction_type'],
        header_rows=2,
        for_tag='{% tr for ap in affiliated_parties %}',
        endfor_tag='{% tr endfor %}',
    )
except Exception as e:
    print(f"  T19 warning: {e}")

print("Converting table T20 (Appendix 3 — Comparable Companies) to Jinja loop...")
try:
    convert_table_to_loop(
        tables[20],
        loop_var='cc',
        fields=['loop.index', 'cc.name', 'cc.description'],
        header_rows=2,
        for_tag='{% tr for cc in comparable_companies %}',
        endfor_tag='{% tr endfor %}',
    )
except Exception as e:
    print(f"  T20 warning: {e}")


# ═════════════════════════════════════════════════════════════════════════════
# 5. IQR TABLE (T16) — simple value replacements in cells
# ═════════════════════════════════════════════════════════════════════════════
print("Templating T16 (IQR table) values...")
try:
    t16 = tables[16]
    for row in t16.rows:
        for cell in row.cells:
            txt = cell.text.strip()
            # Replace numeric values that look like ROS percentages in IQR rows
            for para in cell.paragraphs:
                full = para_full_text(para)
                # Q1 / median / Q3 / IQR rows
                if "Min" in full or "1st Quartile" in full or "Lower Quartile" in full:
                    if any(c.isdigit() for c in full.split()[-1] if c not in (',', '.')):
                        pass  # leave — complex formatting
except Exception as e:
    print(f"  T16 warning: {e}")


# ═════════════════════════════════════════════════════════════════════════════
# 6. ROS CALCULATION TABLE (T17) — template the PT SP rows
# ═════════════════════════════════════════════════════════════════════════════
print("Templating T17 (ROS Calculation table) values...")
try:
    t17 = tables[17]
    # Row 1: Sales - replace the 3 financial year values
    sales_row = t17.rows[2]  # Sales [a]
    cells_17 = sales_row.cells
    if len(cells_17) >= 4:
        set_cell_text(cells_17[1], '{{ financial_data.sales }}')
        set_cell_text(cells_17[2], '{{ financial_data_prior.sales }}')
        set_cell_text(cells_17[3], '{{ financial_data_prior2.sales }}')

    # Operating Income row
    op_row = t17.rows[3]
    cells_op = op_row.cells
    if len(cells_op) >= 4:
        set_cell_text(cells_op[1], '{{ financial_data.operating_profit }}')
        set_cell_text(cells_op[2], '{{ financial_data_prior.operating_profit }}')
        set_cell_text(cells_op[3], '{{ financial_data_prior2.operating_profit }}')

    # ROS % row
    ros_row = t17.rows[4]
    cells_ros = ros_row.cells
    if len(cells_ros) >= 4:
        set_cell_text(cells_ros[1], '{{ ros_current }}%')
        set_cell_text(cells_ros[2], '{{ ros_prior }}%')
        set_cell_text(cells_ros[3], '{{ ros_prior2 }}%')

    # Weighted average row
    wa_row = t17.rows[5]
    cells_wa = wa_row.cells
    if len(cells_wa) >= 4:
        set_cell_text(cells_wa[1], '{{ tested_party_ratio }}%')

    # Header year columns
    hdr_row = t17.rows[1]
    cells_hdr = hdr_row.cells
    if len(cells_hdr) >= 4:
        set_cell_text(cells_hdr[1], 'Total in FY {{ fiscal_year }}')
        set_cell_text(cells_hdr[2], 'Total in FY {{ fiscal_year|int - 1 }}')
        set_cell_text(cells_hdr[3], 'Total in FY {{ fiscal_year|int - 2 }}')
except Exception as e:
    print(f"  T17 warning: {e}")


# ═════════════════════════════════════════════════════════════════════════════
# 7. MANAGEMENT TABLE (T03) — has grouped sections (Commissioner, Directors)
#    Keep caption row [0], replace grouped header labels, template data rows
# ═════════════════════════════════════════════════════════════════════════════
print("Templating T03 (Management table)...")
try:
    t03 = tables[3]
    # Row 0: caption → already handled by global replace
    # Row 1: "Commissioners" group header → keep static
    # Rows 2..N: individual names → template
    # Find first data row after group header
    for r_idx, row in enumerate(t03.rows):
        cells = row.cells
        if len(cells) >= 2:
            pos_text = cells[0].text.strip()
            name_text = cells[1].text.strip()
            # If it looks like a real person entry (not a header)
            if pos_text and name_text and pos_text not in ('', 'Commissioners', 'Directors',
                                                            'Table 3.2 Management PT SP in FY 2024',
                                                            'Table 3.2 Management {{ company_short }} in FY {{ fiscal_year }}'):
                # We can't do a simple for-loop here because the table has sub-groups.
                # Instead, template with indexed access.
                # Leave as-is but replace specific names:
                for mgmt_field in [('Commissioner', 'commissioners[0].name'),
                                   ('President Director', 'directors[0].name'),
                                   ('Directors', 'directors')]:
                    pass
    # Simpler: replace all name values with their jinja equivalents
    # The template table keeps real names for now; implementor replaces manually in Word
    # because management structure varies per client
    print("  T03: management table - individual cells left for manual edit in Word (variable structure per client)")
except Exception as e:
    print(f"  T03 warning: {e}")


# ═════════════════════════════════════════════════════════════════════════════
# 8. SAVE
# ═════════════════════════════════════════════════════════════════════════════
print(f"\nSaving to {DST}...")
doc.save(DST)
print("Done! Open the file in Word to verify.")
print()
print("Next step: install docxtpl and render with:")
print("  from docxtpl import DocxTemplate")
print("  tpl = DocxTemplate('docs_template/TP_Local_File_jinja2.docx')")
print("  tpl.render(context)")
print("  tpl.save('output/result.docx')")
