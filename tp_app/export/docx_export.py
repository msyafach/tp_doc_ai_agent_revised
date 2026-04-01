"""
Export module: generates the Transfer Pricing Local File as a .docx using python-docx.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def update_toc_with_word(docx_path: str) -> tuple[bool, str]:
    """
    Open the generated DOCX in Word (invisible), update all fields + TOC + List of Tables,
    save, and close. This gives the TOC and caption numbering correct page numbers.

    Returns (True, "") on success, or (False, error_message) on failure.
    Requires Microsoft Word to be installed (Windows only).
    """
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        abs_path = os.path.abspath(docx_path)
        doc = word.Documents.Open(abs_path)

        # Update all fields (SEQ, PAGE, NUMPAGES, TOC, List of Tables)
        doc.Fields.Update()
        for toc in doc.TablesOfContents:
            toc.Update()
        for tof in doc.TablesOfFigures:
            tof.Update()

        doc.Save()
        doc.Close(False)
        word.Quit()

        pythoncom.CoUninitialize()
        return True, ""

    except Exception as exc:
        return False, str(exc)


from templates.sections import (
    GLOSSARY,
    STATEMENT_LETTER,
    TP_REGULATIONS_INTRO,
    TP_DOCUMENTATION_TYPES,
    MF_LF_THRESHOLD_TABLE,
    TP_DEADLINE_TEXT,
    SPECIAL_RELATIONSHIP,
    SANCTIONS_TEXT,
    FUNCTIONAL_ANALYSIS_INTRO,
    SEARCH_CRITERIA_DESCRIPTIONS,
    REFERENCES,
    METHOD_DESCRIPTIONS,
    COMPARABILITY_ANALYSIS_INTRO,
    INTERNAL_EXTERNAL_COMPARABLE,
    AFFILIATED_PARTIES_CRITERIA_INTRO,
    AFFILIATED_PARTIES_CRITERIA,
)


# ─── Style helpers ────────────────────────────────────────────────────────────

def insert_toc(doc) -> None:
    """
    Insert a Word TOC field (levels 1-3) that auto-renders when the document
    is opened in Word. Press Ctrl+A → F9, or right-click → Update Field.
    """
    _insert_field(doc, ' TOC \\o "1-3" \\h \\z \\u ',
                  placeholder='[Right-click > Update Field to generate Table of Contents]')


def insert_list_of_tables(doc) -> None:
    """Insert a List of Tables field that collects all Caption-style paragraphs."""
    _insert_field(doc, ' TOC \\h \\z \\c "Table" ',
                  placeholder='[Right-click > Update Field to generate List of Tables]')


def _insert_field(doc, instr: str, placeholder: str = '') -> None:
    """Generic helper: inserts a Word field instruction into a new paragraph."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    fld_begin.set(qn('w:dirty'), 'true')
    run._r.append(fld_begin)

    instr_elem = OxmlElement('w:instrText')
    instr_elem.set(qn('xml:space'), 'preserve')
    instr_elem.text = instr
    run._r.append(instr_elem)

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)

    if placeholder:
        ph_run = paragraph.add_run(placeholder)
        ph_run.italic = True
        ph_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)


def add_table_caption(doc, label: str, source: str = '') -> None:
    """
    Add an auto-numbered table caption using Word's Caption style and SEQ field.
    Format: "Table [auto-number] [label]"
    The List of Tables field collects all paragraphs in Caption style automatically.

    Args:
        doc:    python-docx Document
        label:  descriptive title after the number, e.g. "Ownership Structure PT X in FY 2024"
        source: optional source note printed in italic below the caption
    """
    # Ensure Caption style exists (built-in in Word, may not be in python-docx's default)
    try:
        cap_style = doc.styles['Caption']
    except KeyError:
        cap_style = doc.styles.add_style('Caption', 1)   # 1 = paragraph style
        cap_style.base_style = doc.styles['Normal']
        cap_style.font.bold = True
        cap_style.font.italic = True
        cap_style.font.size = Pt(10)

    p = doc.add_paragraph(style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Literal "Table "
    p.add_run('Table ')

    # Auto-number field: SEQ Table \* ARABIC
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' SEQ Table \\* ARABIC '
    run._r.append(instr)

    fld_end_elem = OxmlElement('w:fldChar')
    fld_end_elem.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end_elem)

    # Title text
    p.add_run(f' {label}')

    # Optional source note
    if source:
        src_p = doc.add_paragraph()
        src_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = src_p.add_run(f'Source: {source}')
        r.italic = True
        r.font.size = Pt(9)


def set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_formatted_paragraph(doc, text, style=None, bold=False, italic=False,
                            font_size=None, alignment=None, space_after=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    if style:
        p.style = style
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


import re

def add_md_paragraph(doc, text: str, style=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     space_after=None) -> None:
    """
    Add a paragraph that parses **bold** markdown into actual Word bold runs.
    All paragraphs are justified by default.
    """
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)

    # Tokenise on **...**
    tokens = re.split(r'(\*\*.*?\*\*)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = p.add_run(token[2:-2])
            run.bold = True
        else:
            p.add_run(token)


def add_multiline_text(doc, text, style=None):
    """Add text that may contain multiple paragraphs, parsing **bold** markdown."""
    paragraphs = text.strip().split('\n\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        # Bullet points
        if para_text.startswith('- ') or para_text.startswith('• '):
            lines = para_text.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('- ') or line.startswith('• '):
                    line = line[2:]
                    add_md_paragraph(doc, line, style='List Bullet',
                                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
                elif line:
                    add_md_paragraph(doc, line, style=style)
        elif para_text.startswith(('1.', '2.', '3.', '4.', '5.')):
            lines = para_text.split('\n')
            for line in lines:
                line = line.strip()
                if line and line[0].isdigit() and '.' in line[:3]:
                    line = line[line.index('.')+1:].strip()
                    add_md_paragraph(doc, line, style='List Number',
                                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
                elif line:
                    add_md_paragraph(doc, line, style=style)
        else:
            # Regular paragraph — join single newlines
            cleaned = para_text.replace('\n', ' ')
            add_md_paragraph(doc, cleaned, style=style)


def _add_section_sources(doc, sources: list, start_num: int = 1) -> None:
    """
    Add footnote-style source references after a section's body text.

    Renders a thin separator line (top-border paragraph, matching Word's
    footnote separator) followed by numbered URL entries in 8pt Calibri italic.

    Args:
        doc:       python-docx Document object.
        sources:   List of URL strings (from Tavily search results).
        start_num: Starting footnote number (use len(prev_sources)+1 for
                   consecutive numbering across sections).
    """
    if not sources:
        return

    # ── Separator line (thin top border, like Word's footnote separator) ──
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(6)
    sep.paragraph_format.space_after  = Pt(2)
    sep_pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '6')        # 0.75pt line
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '000000')
    pBdr.append(top)
    sep_pPr.append(pBdr)

    # ── Numbered source entries ──
    for i, url in enumerate(sources, start_num):
        entry = doc.add_paragraph()
        entry.paragraph_format.space_before = Pt(0)
        entry.paragraph_format.space_after  = Pt(1)
        run = entry.add_run(f"[{i}] {url}")
        run.font.name   = 'Calibri'
        run.font.size   = Pt(8)
        run.font.italic = True


def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(header_cells[i], "D5E8F0")
    
    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    
    # Set column widths
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Inches(width)
    
    return table


# ─── Main export function ────────────────────────────────────────────────────

def generate_tp_document(state: dict, output_path: str) -> str:
    """
    Generate the complete Transfer Pricing Local File document.
    
    Args:
        state: Complete state dictionary with all data (manual + agent-generated)
        output_path: Path to save the .docx file
    
    Returns:
        Path to the generated document
    """
    doc = Document()
    
    # ─── Document styles ──────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ─── Page numbers in footer ───────────────────────────────────
    def _make_field_run(paragraph, fld_instr: str):
        """Helper: append a simple Word field to a paragraph run."""
        run = paragraph.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_begin)

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f' {fld_instr} '
        run._r.append(instr)

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_end)

    def _add_page_numbers_to_section(section):
        """Add 'Page X of Y' to the footer of a given section."""
        footer = section.footer
        footer.is_linked_to_previous = False
        # Clear any existing paragraphs
        for para in footer.paragraphs:
            p = para._element
            p.getparent().remove(p)

        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("Page ")
        _make_field_run(para, "PAGE")
        para.add_run(" of ")
        _make_field_run(para, "NUMPAGES")
        # Style the footer text
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Apply to the default (first) section — additional sections inherit it
    _add_page_numbers_to_section(doc.sections[0])


    
    # ─── Extract state data ───────────────────────────────────────
    company_name = state.get("company_name", "[Company Name]")
    company_short = state.get("company_short_name", "[Short Name]")
    fiscal_year = state.get("fiscal_year", "[FY]")
    company_address = state.get("company_address", "[Address]")
    
    # ════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()
    
    add_formatted_paragraph(doc, "TRANSFER PRICING DOCUMENTATION", bold=True,
                           font_size=20, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "LOCAL FILE", bold=True,
                           font_size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_formatted_paragraph(doc, company_name.upper(), bold=True,
                           font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_formatted_paragraph(doc, f"FISCAL YEAR {fiscal_year}", bold=True,
                           font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # GLOSSARY AND ABBREVIATIONS
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("GLOSSARY AND ABBREVIATIONS", level=1)
    
    # Add company-specific abbreviations
    extra_glossary = {}
    if company_short and company_name:
        extra_glossary[company_short] = company_name
    
    all_glossary = {**GLOSSARY, **extra_glossary}
    
    headers = ["Term", "Definition"]
    rows = [[k, v] for k, v in sorted(all_glossary.items())]
    create_table(doc, headers, rows, col_widths=[2.0, 4.5])
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # STATEMENT LETTER
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("STATEMENT LETTER ON SCOPE AND LIMITATION WORK", level=1)
    
    statement = STATEMENT_LETTER.format(
        company_name=company_name,
        company_short_name=company_short,
        fiscal_year=fiscal_year,
    )
    add_multiline_text(doc, statement)
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("TABLE OF CONTENTS", level=1)
    insert_toc(doc)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # LIST OF TABLES
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("LIST OF TABLES", level=1)
    insert_list_of_tables(doc)

    doc.add_page_break()
    # ════════════════════════════════════════════════════════════════
    # CHAPTER 1: EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("EXECUTIVE SUMMARY", level=1)
    
    exec_summary = state.get("executive_summary", "")
    if exec_summary:
        # Parse the sections from the generated text
        add_multiline_text(doc, exec_summary)
    else:
        doc.add_paragraph("[Executive Summary will be generated by the agent]")
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # CHAPTER 2: TRANSFER PRICING REGULATIONS
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("TRANSFER PRICING REGULATIONS", level=1)
    
    doc.add_heading("Transfer Pricing Regulations in Indonesia", level=2)
    add_multiline_text(doc, TP_REGULATIONS_INTRO)
    add_multiline_text(doc, TP_DOCUMENTATION_TYPES)
    
    # Threshold table
    doc.add_paragraph()
    add_table_caption(doc,
        "Requirement of Master File and Local File Preparation",
        source="Article 16 paragraph (3) in PMK-172 of 2023")
    headers = ["No.", "Criteria", "Threshold"]
    rows = [[t["no"], t["criteria"], t["threshold"]] for t in MF_LF_THRESHOLD_TABLE]
    create_table(doc, headers, rows, col_widths=[0.5, 3.0, 3.0])
    
    doc.add_paragraph()
    add_multiline_text(doc, TP_DEADLINE_TEXT)
    add_multiline_text(doc, SANCTIONS_TEXT)
    
    doc.add_heading("Special Relationship", level=2)
    add_multiline_text(doc, SPECIAL_RELATIONSHIP)
    
    # TP Methods overview
    doc.add_heading("Transfer Pricing Methods", level=2)
    for method_name, method_desc in METHOD_DESCRIPTIONS.items():
        doc.add_heading(f"{method_name} Method", level=3)
        add_multiline_text(doc, method_desc)
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # CHAPTER 3: IDENTITY AND BUSINESS ACTIVITIES
    # ════════════════════════════════════════════════════════════════
    doc.add_heading(f"IDENTITY AND BUSINESS ACTIVITIES OF {company_short}", level=1)
    
    # 3.1 Management Structure
    doc.add_heading("Explanation of Management Structure, Organization Structure, and Information Regarding Related Parties", level=2)
    
    # Ownership Structure
    doc.add_heading(f"Ownership Structure and Management of {company_short}", level=3)
    
    shareholders = state.get("shareholders", [])
    if shareholders:
        add_table_caption(doc,
            f"Ownership Structure {company_short} in FY {fiscal_year}",
            source=f"Management {company_short}, 31 December {fiscal_year}")
        headers = ["Shareholders", "Number of Shares", "Total Paid in Capital (IDR)", "Ownership (%)"]
        rows = [[s.get("name",""), s.get("shares",""), s.get("capital",""), s.get("percentage","")]
                for s in shareholders]
        create_table(doc, headers, rows, col_widths=[2.5, 1.5, 2.0, 1.0])
    
    # Management
    management = state.get("management", [])
    if management:
        doc.add_paragraph()
        add_table_caption(doc,
            f"Management {company_short} in FY {fiscal_year}",
            source=f"Management {company_short}, 31 December {fiscal_year}")
        headers = ["Position", "Name"]
        rows = [[m.get("position",""), m.get("name","")] for m in management]
        create_table(doc, headers, rows, col_widths=[3.0, 3.5])
    
    # Organization Structure
    doc.add_heading("Organization Structure", level=3)
    employee_count = state.get("employee_count", "[number]")
    doc.add_paragraph(
        f"In the Fiscal Year {fiscal_year}, {company_short} has a total of permanent employees of {employee_count}. "
        f"Please refer to Appendix 1 regarding the details of the Organization Structure of {company_short}."
    )
    
    # Affiliated Parties
    doc.add_heading(f"Information regarding Parties Having Special Relationships with {company_short}", level=3)
    parent_group = state.get("parent_group", "[Group]")
    doc.add_paragraph(
        f"The parties with a special relationship with {company_short} include its shareholders and other companies "
        f"under {parent_group}. Among these related parties are parties who engage in transactions with {company_short}. "
        f"Details of these companies can be found in Appendix 2."
    )
    
    # 3.2 Business Activities
    doc.add_heading(f"Explanation of Business Activities, Business Strategies of {company_short}", level=2)
    
    doc.add_heading("Business Activities and Operational Aspects", level=3)
    business_desc = state.get("business_activities_description", "")
    if business_desc:
        add_multiline_text(doc, business_desc)
    
    # Products table — styled: dark caption, green header, blue name cells, images
    products = state.get("products", [])
    product_images = state.get("product_images", {})

    if products and any(p.get("name") for p in products):
        # ── Row 0: dark caption using add_table_caption *before* the table ──
        # We add the caption paragraph before the table, then build the table
        add_table_caption(doc,
            f"Details of Products {company_short}",
            source=f"Management {company_short}")

        num_products = len(products)
        # Table: header row + product rows (caption is now a separate paragraph above)
        table = doc.add_table(rows=1 + num_products, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # ── Row 0: green header "Products" ──
        hdr_cell = table.cell(0, 0)
        hdr_cell.merge(table.cell(0, 1))
        hdr_cell.text = ""
        p = hdr_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Products")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cell, "4CAF50")  # green

        # ── Product rows ──
        import tempfile as _tmpfile
        for idx, prod in enumerate(products):
            row_idx = 1 + idx
            name = prod.get("name", "")

            # Left cell — sky blue, white bold text
            name_cell = table.cell(row_idx, 0)
            name_cell.text = ""
            p = name_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(name_cell, "29ABE2")  # sky blue

            # Vertically centre the name cell
            tc_pr = name_cell._element.get_or_add_tcPr()
            v_align = OxmlElement('w:vAlign')
            v_align.set(qn('w:val'), 'center')
            tc_pr.append(v_align)

            # Right cell — product image (if uploaded)
            img_cell = table.cell(row_idx, 1)
            img_cell.text = ""
            img_bytes = product_images.get(idx) or product_images.get(str(idx))
            if img_bytes:
                p = img_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                tmp = _tmpfile.NamedTemporaryFile(suffix=".png", delete=False)
                try:
                    tmp.write(img_bytes)
                    tmp.close()
                    run.add_picture(tmp.name, width=Inches(2.0))
                finally:
                    try:
                        os.unlink(tmp.name)
                    except Exception:
                        pass

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(3.0)
            row.cells[1].width = Inches(3.5)

        add_formatted_paragraph(doc, f"Source: Management {company_short}",
                               italic=True, font_size=9)

    
    # Business Strategy
    doc.add_heading(f"Business Strategy {company_short}", level=3)
    biz_strategy = state.get("business_strategy", "")
    if biz_strategy:
        add_multiline_text(doc, biz_strategy)
    
    # Business Restructuring
    doc.add_heading("Business Restructuring or Transfer of Intangible Assets", level=3)
    restructuring = state.get("business_restructuring", 
        f"In the Fiscal Year {fiscal_year}, {company_short} is not involved or affected by any "
        f"business restructuring or transfer of intangible assets within the Business Group.")
    doc.add_paragraph(restructuring)
    
    # 3.3 Business Environment (Agent-generated)
    doc.add_heading("Overview of the Business Environment Including List of Main Competitors", level=2)
    biz_env = state.get("business_environment_overview", "")
    if biz_env:
        add_multiline_text(doc, biz_env)
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # CHAPTER 4: AFFILIATED TRANSACTIONS & FUNCTIONAL ANALYSIS
    # ════════════════════════════════════════════════════════════════
    doc.add_heading(f"INFORMATION ON AFFILIATED AND INDEPENDENT TRANSACTIONS CONDUCTED BY {company_short} AND FUNCTIONAL ANALYSIS", level=1)
    
    # 4.1 Transaction Scheme
    doc.add_heading("Transaction Scheme, Pricing Policy, and Transaction Details", level=2)
    
    # Affiliated parties list
    doc.add_heading("List of Affiliated Parties", level=3)

    # Template: Article 18 paragraph (4) criteria
    doc.add_paragraph(AFFILIATED_PARTIES_CRITERIA_INTRO)
    for i, criteria_text in enumerate(AFFILIATED_PARTIES_CRITERIA):
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{chr(97 + i)}) {criteria_text}")

    # Closing sentence referencing Appendix 2
    closing = doc.add_paragraph(
        f"Based on the criteria for related party according to Article 18 paragraph (4) of the "
        f"Income Tax Law, the list of {company_short} affiliated party can be seen in the table below."
    )

    # Affiliated parties table
    affiliated = state.get("affiliated_parties", [])
    if affiliated:
        add_table_caption(doc, f"List of Affiliated Parties of {company_short}")
        headers = ["No.", "Company Name", "Country", "Relationship", "Transaction Type"]
        rows = [[str(i+1), a.get("name",""), a.get("country",""), a.get("relationship",""), a.get("transaction_type","")]
                for i, a in enumerate(affiliated)]
        create_table(doc, headers, rows, col_widths=[0.4, 2.0, 1.2, 1.5, 1.5])
    
    # Transaction details
    doc.add_heading("Transaction Details", level=3)
    tx_details = state.get("transaction_details_text", "")
    if tx_details:
        add_multiline_text(doc, tx_details)
    
    # Pricing policy
    doc.add_heading("Pricing Policy", level=3)
    pricing = state.get("pricing_policy", "")
    if pricing:
        add_multiline_text(doc, pricing)

    # 4.1.2.4 Copy of Agreement/Contract Related to Transactions of Significant Value
    doc.add_heading(
        "Copy of Agreement/Contract Related to Transactions of Significant Value",
        level=3,
    )
    doc.add_paragraph(
        "A copy of the agreement/contract for the affiliated transaction is attached in Appendix 2."
    )

    # 4.2 Functional Analysis (Agent-generated)

    doc.add_heading("Functional Analysis", level=2)
    
    doc.add_heading("Supply Chain Management", level=3)
    add_multiline_text(doc, FUNCTIONAL_ANALYSIS_INTRO[:500])  # Standard intro
    
    func_analysis = state.get("functional_analysis_narrative", "")
    if func_analysis:
        add_multiline_text(doc, func_analysis)
    
    # Business Characterization — left blank for user to fill manually
    doc.add_heading(f"Business Characteristics of {company_short}", level=3)
    # 🟡 Manual — user fills this section in the exported document
    doc.add_paragraph("")   # blank line 1
    doc.add_paragraph("")   # blank line 2
    doc.add_paragraph("")   # blank line 3

    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # CHAPTER 5: APPLICATION OF ARM'S LENGTH PRINCIPLE
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("APPLICATION OF ARM'S LENGTH PRINCIPLE OF THE AFFILIATED PARTY TRANSACTIONS", level=1)

    # Static template intro paragraph (🟢 Template — not user-editable)
    doc.add_paragraph(
        "This section explains the steps that need to be taken in applying the arm\u2019s length "
        "principle such as industry analysis, analysis of transaction conditions, comparability "
        "analysis, selection of the most suitable transfer pricing method, advantages and "
        "disadvantages of each method, profit level indicators, search, and the interquartile "
        "range of comparable companies."
    )

    # 5.1 Industry Analysis (Agent-generated)
    doc.add_heading("Industry Analysis", level=2)

    # Derive industry name from products list (same logic as docx_template_export.py)
    _products = state.get("products", [])
    _industry_name = (
        state.get("industry_name", "")
        or ", ".join(p.get("name", "") for p in _products if p.get("name"))
        or "Industry"
    )[:60].strip()

    doc.add_heading(f"Global {_industry_name} Industry Analysis", level=3)
    global_industry = state.get("industry_analysis_global", "")
    if global_industry:
        add_multiline_text(doc, global_industry)
    # Footnote-style sources for global industry analysis
    _global_sources = state.get("industry_global_sources", [])
    _add_section_sources(doc, _global_sources, start_num=1)

    doc.add_heading(f"Indonesian {_industry_name} Industry Analysis", level=3)
    indo_industry = state.get("industry_analysis_indonesia", "")
    if indo_industry:
        add_multiline_text(doc, indo_industry)
    # Footnote-style sources for Indonesian industry analysis
    _indo_sources = state.get("industry_indonesia_sources", [])
    _add_section_sources(doc, _indo_sources, start_num=len(_global_sources) + 1)

    # 5.1.2.1 Efficiency and Excellence Levels of Company Locations (Agent-generated)
    doc.add_heading("Efficiency and Excellence Levels of Company Locations", level=3)
    location_analysis = state.get("company_location_analysis", "")
    if location_analysis:
        add_multiline_text(doc, location_analysis)
    _loc_sources = state.get("company_location_sources", [])
    _add_section_sources(
        doc, _loc_sources,
        start_num=len(_global_sources) + len(_indo_sources) + 1,
    )

    # 5.1.2.2 Regulations Affecting the Industry (Agent-generated)
    doc.add_heading("Regulations Affecting the Industry", level=3)
    regulations_text = state.get("industry_regulations_text", "")
    if regulations_text:
        add_multiline_text(doc, regulations_text)
    _reg_sources = state.get("industry_regulations_sources", [])
    _add_section_sources(
        doc, _reg_sources,
        start_num=len(_global_sources) + len(_indo_sources) + len(_loc_sources) + 1,
    )

    # 5.2 Comparability Analysis

    doc.add_heading("Comparability Analysis", level=2)
    comp_analysis = state.get("comparability_analysis_narrative", "")
    if comp_analysis:
        add_multiline_text(doc, comp_analysis)

    # Table 5.1: Comparability Analysis Factors (manually filled)
    _comp_factors = state.get("comparability_factors", [])
    if _comp_factors:
        _company_short = state.get("company_short_name", "Company")
        _fy            = state.get("fiscal_year", "")
        _cap_text      = f"Table 5.1 Comparability Analysis of {_company_short} Affiliated Transactions FY {_fy}"

        # Caption paragraph (table number style)
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(_cap_text)
        cap_run.bold = True
        cap_run.font.size = Pt(10)

        table = doc.add_table(rows=1 + len(_comp_factors), cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_labels = ["No.", "Comparability Factors", "Description"]
        hdr_colors = ["1F5C00", "1F5C00", "1F5C00"]   # dark green header
        for ci, (label, color) in enumerate(zip(hdr_labels, hdr_colors)):
            cell = hdr_cells[ci]
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), color)
            cell._tc.get_or_add_tcPr().append(shading)
            run = cell.paragraphs[0].add_run(label)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0x00)  # yellow text (matches template)
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set column widths: No. | Factor | Description
        col_widths = [Cm(1.0), Cm(4.5), Cm(10.5)]
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = w

        # Data rows (alternating light/white)
        _row_colors = ["F2F2F2", "FFFFFF"]
        for ri, row_data in enumerate(_comp_factors):
            row = table.rows[ri + 1]
            bg = _row_colors[ri % 2]

            for ci, val in enumerate([str(ri + 1), row_data.get("factor", ""), row_data.get("description", "")]):
                cell = row.cells[ci]
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), bg)
                cell._tc.get_or_add_tcPr().append(shading)
                para = cell.paragraphs[0]
                run  = para.add_run(val)
                run.font.size = Pt(9)
                if ci == 1:  # Factor column — bold
                    run.bold = True
                if ci == 0:  # No. column — center align
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacing after table


    # ─ 5.3 intro: Comparability Analysis explanation + 5 factors (a–e) ─────────
    add_multiline_text(doc, COMPARABILITY_ANALYSIS_INTRO)

    # 5.3.1 Use of Internal and External Comparable
    doc.add_heading("Use of Internal and External Comparable in Transfer Pricing Analysis", level=3)
    add_multiline_text(doc, INTERNAL_EXTERNAL_COMPARABLE)

    # 5.4 Selection of TP Methods (Agent-generated)
    doc.add_heading("Selection of Transfer Pricing Methods", level=2)
    method_just = state.get("method_selection_justification", "")
    if method_just:
        add_multiline_text(doc, method_just)
    
    # PLI Selection
    doc.add_heading("Selection of Profit Level Indicators (PLI)", level=3)
    pli_rationale = state.get("pli_selection_rationale", "")
    if pli_rationale:
        add_multiline_text(doc, pli_rationale)
    
    # 5.4 Search Summary
    doc.add_heading("Application of the Accepted Method", level=2)
    doc.add_heading("Search Summary", level=3)
    
    search_criteria = state.get("search_criteria_results", [])
    if search_criteria:
        for desc_key, desc_text in SEARCH_CRITERIA_DESCRIPTIONS.items():
            doc.add_paragraph(desc_text)

        add_table_caption(doc, f"Comparable Company Search Summary")
        headers = ["No.", "Criteria", "Result"]
        rows = [[str(s.get("step","")), s.get("criteria",""), s.get("result_count","")]
                for s in search_criteria]
        create_table(doc, headers, rows, col_widths=[0.5, 4.0, 1.5])
    
    # Comparable companies
    doc.add_heading("Selected Comparable Companies", level=3)
    comparables = state.get("comparable_companies", [])
    if comparables:
        add_table_caption(doc, f"Selected Comparable Companies")
        headers = ["No.", "Company Name", "Country", "Description"]
        rows = [[str(i+1), c.get("name",""), c.get("country",""), c.get("description","")[:100]]
                for i, c in enumerate(comparables)]
        create_table(doc, headers, rows, col_widths=[0.4, 2.0, 1.0, 3.0])
    
    # Results
    doc.add_heading("Results of Method Application", level=3)
    selected_method = state.get("selected_method", "TNMM")
    selected_pli = state.get("selected_pli", "ROS")
    quartile = state.get("quartile_range", {})
    tested_ratio = state.get("tested_party_ratio", 0)
    analysis_period = state.get("analysis_period", "")
    
    if quartile:
        add_table_caption(doc,
            f"Results of {selected_method} Method Application — {company_short} {selected_pli}")
        headers = ["Metric", "Value"]
        rows = [
            ["1st Quartile", f"{quartile.get('q1', 0):.2f}%"],
            ["Median", f"{quartile.get('median', 0):.2f}%"],
            ["3rd Quartile", f"{quartile.get('q3', 0):.2f}%"],
            [f"{company_short}'s Weighted Average {selected_pli}", f"{tested_ratio:.2f}%"],
        ]
        create_table(doc, headers, rows, col_widths=[3.5, 2.0])
    
    # 5.5 Conclusion (Agent-generated)
    doc.add_heading("Conclusion", level=2)
    conclusion = state.get("conclusion_text", "")
    if conclusion:
        add_multiline_text(doc, conclusion)
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # CHAPTER 6: FINANCIAL INFORMATION
    # ════════════════════════════════════════════════════════════════
    doc.add_heading(f"{company_short} FINANCIAL INFORMATION", level=1)
    
    doc.add_heading(f"Overview of Profit/Loss of {company_short}", level=2)
    
    fin_data = state.get("financial_data", {})
    fin_prior = state.get("financial_data_prior", {})
    
    if fin_data:
        add_table_caption(
            doc,
            f"Profit/Loss Summary of {company_short} for FY {fiscal_year}",
            source=f"Audited Financial Statement {company_short} FY {fiscal_year}"
        )
        
        pl_items = ["sales", "cogs", "gross_profit", "gross_margin_pct", 
                    "operating_expenses", "operating_profit", "financial_income",
                    "other_income", "other_expense", "income_before_tax", 
                    "income_tax", "net_income"]
        pl_labels = {
            "sales": "Sales",
            "cogs": "Cost of Goods Sold",
            "gross_profit": "Gross Profit",
            "gross_margin_pct": "Gross Profit Margin (%)",
            "operating_expenses": "Operational Expenses",
            "operating_profit": "Operating Profit",
            "financial_income": "Financial Income (Expenses)",
            "other_income": "Other Income",
            "other_expense": "Other Expense",
            "income_before_tax": "Income (Loss) Before Tax",
            "income_tax": "Income Tax Expenses",
            "net_income": "Net (Loss) Income",
        }
        
        headers = ["Profit/Loss Account", fiscal_year, str(int(fiscal_year)-1) if fiscal_year.isdigit() else "Prior Year"]
        rows = []
        for item in pl_items:
            current = fin_data.get(item, "")
            prior = fin_prior.get(item, "")
            rows.append([pl_labels.get(item, item), str(current), str(prior)])
        
        create_table(doc, headers, rows, col_widths=[2.5, 2.0, 2.0])
        add_formatted_paragraph(doc, f"Source: Management {company_short}, 31 December {fiscal_year}",
                               italic=True, font_size=9)

        # P/L overview bullet points (AI-generated)
        pl_overview = state.get("pl_overview_text", "")
        if pl_overview:
            doc.add_paragraph()  # spacing
            add_multiline_text(doc, pl_overview)
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # CHAPTER 7: NON-FINANCIAL EVENTS
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("NON-FINANCIAL EVENTS/OCCURRENCES/FACTS THAT AFFECT PRICING OR PROFIT LEVELS", level=1)
    
    non_fin = state.get("non_financial_events", 
        f"In FY {fiscal_year}, there were no significant non-financial events that affected the pricing or profit levels of {company_short}.")
    add_multiline_text(doc, non_fin)
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # REFERENCES
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("REFERENCES", level=1)
    for i, ref in enumerate(REFERENCES, 1):
        doc.add_paragraph(f"[{i}] {ref}")
    
    # Add industry analysis sources
    global_sources = state.get("industry_global_sources", [])
    indo_sources = state.get("industry_indonesia_sources", [])
    all_sources = global_sources + indo_sources
    for i, src in enumerate(all_sources, len(REFERENCES) + 1):
        doc.add_paragraph(f"[{i}] {src}")
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # APPENDICES
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("APPENDIX 1: ORGANIZATION STRUCTURE", level=1)

    org_desc = state.get("org_structure_description", "")
    org_depts = state.get("org_structure_departments", [])
    org_image_bytes = state.get("org_structure_image", None)

    if org_desc:
        add_multiline_text(doc, org_desc)

    if org_depts and any(d.get("name") for d in org_depts):
        add_table_caption(doc, f"Organization Structure of {company_short} — FY {fiscal_year}")
        headers = ["Department / Unit", "Head / PIC", "Number of Employees"]
        rows = [[
            d.get("name", ""),
            d.get("head", ""),
            d.get("employees", ""),
        ] for d in org_depts if d.get("name")]
        create_table(doc, headers, rows, col_widths=[3.0, 2.5, 1.5])

    if org_image_bytes:
        import tempfile as _tmpfile2
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        tmp2 = _tmpfile2.NamedTemporaryFile(suffix=".png", delete=False)
        try:
            tmp2.write(org_image_bytes)
            tmp2.close()
            run.add_picture(tmp2.name, width=Inches(6.0))
        finally:
            try:
                os.unlink(tmp2.name)
            except Exception:
                pass
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Figure 1: Organization Chart of {company_short}")
        r.italic = True
        r.font.size = Pt(9)
    elif not org_desc and not (org_depts and any(d.get("name") for d in org_depts)):
        doc.add_paragraph("[Insert organization structure chart — use Step 4 to upload]")

    
    doc.add_page_break()
    doc.add_heading("APPENDIX 2: LIST OF AFFILIATED PARTIES", level=1)
    if affiliated:
        headers = ["No.", "Company", "Country", "Relationship", "Transactions"]
        rows = [[str(i+1), a.get("name",""), a.get("country",""), 
                 a.get("relationship",""), a.get("transaction_type","")] 
                for i, a in enumerate(affiliated)]
        create_table(doc, headers, rows, col_widths=[0.4, 2.0, 1.0, 1.5, 1.5])
    
    doc.add_page_break()
    doc.add_heading("APPENDIX 3: OVERVIEW OF COMPARABLE COMPANIES", level=1)
    if comparables:
        for i, c in enumerate(comparables, 1):
            doc.add_heading(f"{i}. {c.get('name', '')}", level=3)
            doc.add_paragraph(f"Country: {c.get('country', '')}")
            doc.add_paragraph(c.get('description', ''))
    
    doc.add_page_break()
    doc.add_heading("APPENDIX 4: FINANCIAL STATEMENT OF COMPARABLE COMPANIES", level=1)
    doc.add_paragraph("[Insert comparable company financial data tables]")
    
    doc.add_page_break()
    doc.add_heading("APPENDIX 5: REJECTION MATRIX OF COMPARABLE COMPANIES", level=1)
    doc.add_paragraph("[Insert rejection matrix]")
    
    # ─── Save ─────────────────────────────────────────────────────
    doc.save(output_path)
    return output_path
